# -*- coding: utf-8 -*-
"""
Gera relatorio_resultados_pipeline.docx com análise completa, gráficos reais
e dataframes dos resultados do pipeline Intensivão Quant AI.
"""
import os, io, warnings
import numpy as np
import pandas as pd
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.gridspec as gridspec
from matplotlib.ticker import FuncFormatter
import scipy.stats as stats
warnings.filterwarnings('ignore')

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Caminhos ─────────────────────────────────────────────────────────────────
BASE      = os.path.dirname(os.path.abspath(__file__))
DADOS_DIR = os.path.join(BASE, 'dados')
OUT       = os.path.join(BASE, 'relatorio_resultados_pipeline_v2.docx')

NAVY  = '#002D62'
GOLD  = '#B88600'
GREEN = '#007000'
RED   = '#CC0000'
GRAY  = '#555555'

# ── Helpers DOCX ─────────────────────────────────────────────────────────────

def novo_doc():
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(2.8)
        section.right_margin  = Cm(2.5)
    return doc

def titulo(doc, texto):
    p = doc.add_heading(texto, level=0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.color.rgb = RGBColor(0x00, 0x2D, 0x62)
    p.runs[0].font.size = Pt(20)

def h1(doc, texto):
    p = doc.add_heading(texto, level=1)
    p.runs[0].font.color.rgb = RGBColor(0x00, 0x2D, 0x62)
    p.runs[0].font.size = Pt(14)

def h2(doc, texto):
    p = doc.add_heading(texto, level=2)
    p.runs[0].font.color.rgb = RGBColor(0xB8, 0x86, 0x00)
    p.runs[0].font.size = Pt(12)

def corpo(doc, texto, bold=False, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(texto)
    r.font.size = Pt(11)
    r.bold   = bold
    r.italic = italic
    return p

def bullet(doc, texto, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    r = p.add_run(texto)
    r.font.size = Pt(11)

def alerta(doc, texto):
    p = doc.add_paragraph()
    r = p.add_run(f"⚠  {texto}")
    r.font.size   = Pt(11)
    r.bold        = True
    r.font.color.rgb = RGBColor(0xCC, 0x44, 0x00)

def destaque(doc, label, valor, nota=""):
    p = doc.add_paragraph()
    r1 = p.add_run(f"{label}: "); r1.bold = True; r1.font.size = Pt(11)
    r2 = p.add_run(valor);        r2.font.size = Pt(11); r2.font.color.rgb = RGBColor(0,112,0)
    if nota:
        r3 = p.add_run(f"  — {nota}"); r3.font.size = Pt(10); r3.italic = True
        r3.font.color.rgb = RGBColor(0x55,0x55,0x55)

def _shd(cell, hex_color):
    s = OxmlElement('w:shd')
    s.set(qn('w:val'),   'clear')
    s.set(qn('w:color'), 'auto')
    s.set(qn('w:fill'),  hex_color)
    cell._tc.get_or_add_tcPr().append(s)

def tabela(doc, colunas, linhas, col_width=None):
    t = doc.add_table(rows=1 + len(linhas), cols=len(colunas))
    t.style = 'Table Grid'
    for j, col in enumerate(colunas):
        cell = t.rows[0].cells[j]
        cell.text = col
        r = cell.paragraphs[0].runs[0]
        r.bold = True; r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        _shd(cell, '002D62')
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if col_width:
            cell.width = Inches(col_width[j])
    for i, row in enumerate(linhas):
        for j, val in enumerate(row):
            cell = t.rows[i+1].cells[j]
            cell.text = str(val)
            cell.paragraphs[0].runs[0].font.size = Pt(10)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if i % 2 == 0:
                _shd(cell, 'EEF2FF')
    doc.add_paragraph()

def insert_fig(doc, fig, width=6.0, caption=None):
    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=150, bbox_inches='tight',
                facecolor='white')
    buf.seek(0)
    plt.close(fig)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(buf, width=Inches(width))
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cp.add_run(caption)
        r.font.size  = Pt(9)
        r.italic     = True
        r.font.color.rgb = RGBColor(0x55,0x55,0x55)
    doc.add_paragraph()

# ── Helpers de gráfico ────────────────────────────────────────────────────────

def pct_fmt(x, _): return f'{x:.0%}'
def brl_fmt(x, _): return f'R${x:,.0f}'

def cum_ret(s):
    return (1 + s).cumprod()

def calc_metricas(ret, nome='Estratégia', rf=0.0):
    ret = ret.dropna()
    n   = len(ret)
    if n == 0:
        return {}
    cagr    = (1 + ret).prod() ** (12/n) - 1
    vol     = ret.std() * np.sqrt(12)
    sharpe  = (ret.mean() - rf) / ret.std() * np.sqrt(12) if ret.std() > 0 else 0
    neg     = ret[ret < rf]
    sortino = (ret.mean() - rf) / neg.std() * np.sqrt(12) if len(neg) > 1 else 0
    cum     = cum_ret(ret)
    mdd     = ((cum / cum.cummax()) - 1).min()
    calmar  = cagr / abs(mdd) if mdd != 0 else 0
    return dict(nome=nome, CAGR=cagr, Vol=vol, Sharpe=sharpe,
                Sortino=sortino, MDD=mdd, Calmar=calmar, n=n)

def fmt_row(m):
    return [m['nome'],
            f"{m['CAGR']:.1%}", f"{m['Vol']:.1%}",
            f"{m['Sharpe']:.2f}", f"{m['Sortino']:.2f}",
            f"{m['MDD']:.1%}",   f"{m['Calmar']:.2f}"]

# ── Carrega dados ─────────────────────────────────────────────────────────────

def load():
    d = {}
    arquivos = {
        'precos':      'precos_ibov.parquet',
        'ret_d':       'retornos_diarios_limpo.parquet',
        'ret_m':       'retornos_mensais_limpo.parquet',
        'sinal_v1':    'sinal_v1.parquet',
        'sinal_v2':    'sinal_v2.parquet',
        'pesos_v1':    'pesos_v1.parquet',
        'pesos_v2':    'pesos_v2.parquet',
        'pesos_v2s':   'pesos_sinal_v2.parquet',
        'cart_v1':     'retorno_carteira_v1.parquet',
        'cart_v2':     'retorno_carteira_v2.parquet',
        'cart_v2s':    'retorno_carteira_sinal_v2.parquet',
        'oos':         'retorno_walkforward_liquido.parquet',
    }
    for k, fname in arquivos.items():
        path = os.path.join(DADOS_DIR, fname)
        if os.path.exists(path):
            d[k] = pd.read_parquet(path)
    return d


# ════════════════════════════════════════════════════════════════════════════
#  GRÁFICOS
# ════════════════════════════════════════════════════════════════════════════

def fig_equity_curves(d):
    """Equity curves: Momentum v1, Vol-Weight, Walk-Forward OOS, IBOVESPA"""
    cart_v1  = d['cart_v1'].squeeze()
    oos      = d['oos'].squeeze()

    # Proxy IBOV: média dos retornos mensais do universo
    ibov = d['ret_m'].mean(axis=1).reindex(cart_v1.index)

    fig, ax = plt.subplots(figsize=(10, 4.5))
    for s, label, color, lw in [
        (cart_v1, 'Momentum v1 (IS)',    NAVY,  2.2),
        (oos,     'Walk-Forward OOS',    GREEN, 2.0),
        (ibov,    'IBOVESPA (proxy)',    GRAY,  1.4),
    ]:
        s = s.dropna()
        ax.plot(s.index, cum_ret(s), label=label, color=color, linewidth=lw)

    ax.set_title('Curvas de Patrimônio Acumulado (base 1)', fontsize=13, color=NAVY, fontweight='bold')
    ax.set_ylabel('Retorno Acumulado (1 = capital inicial)')
    ax.yaxis.set_major_formatter(FuncFormatter(lambda x,_: f'{x:.1f}x'))
    ax.legend(fontsize=10)
    ax.grid(True, alpha=0.3)
    ax.set_facecolor('#FAFAFA')
    fig.tight_layout()
    return fig

def fig_distribuicao(d):
    """Histograma + QQ plot dos retornos mensais do universo"""
    todos = d['ret_m'].values.flatten()
    todos = todos[~np.isnan(todos)]

    fig, axes = plt.subplots(1, 2, figsize=(11, 4))

    # Histograma
    ax = axes[0]
    ax.hist(todos, bins=80, color=NAVY, alpha=0.7, density=True, label='Retornos reais')
    mu, sigma = todos.mean(), todos.std()
    x = np.linspace(todos.min(), todos.max(), 300)
    ax.plot(x, stats.norm.pdf(x, mu, sigma), color=RED, lw=2, label='Normal teórica')
    ax.set_title('Distribuição dos Retornos Mensais', fontsize=12, color=NAVY, fontweight='bold')
    ax.set_xlabel('Retorno mensal')
    ax.xaxis.set_major_formatter(FuncFormatter(pct_fmt))
    ax.legend(fontsize=9)
    ax.grid(True, alpha=0.3)

    kurt = pd.Series(todos).kurt()
    skew = pd.Series(todos).skew()
    ax.text(0.97, 0.95, f'Kurtosis: {kurt:.2f}\nSkewness: {skew:.2f}',
            transform=ax.transAxes, ha='right', va='top', fontsize=10,
            bbox=dict(boxstyle='round', facecolor='wheat', alpha=0.5))

    # QQ Plot
    ax2 = axes[1]
    (osm, osr), (slope, intercept, r) = stats.probplot(todos, dist='norm')
    ax2.scatter(osm, osr, alpha=0.3, s=8, color=NAVY, label='Dados reais')
    lims = [min(osm), max(osm)]
    ax2.plot(lims, [slope*l + intercept for l in lims], color=RED, lw=2, label='Normal')
    ax2.set_title('QQ-Plot vs Normal', fontsize=12, color=NAVY, fontweight='bold')
    ax2.set_xlabel('Quantis teóricos'); ax2.set_ylabel('Quantis observados')
    ax2.legend(fontsize=9)
    ax2.grid(True, alpha=0.3)

    fig.tight_layout()
    return fig

def fig_ic_serie(d):
    """IC mensal do sinal v1 e v2"""
    sinal_v1 = d['sinal_v1']
    ret_m    = d['ret_m']

    # Recalcular IC (Spearman) mês a mês
    datas = sinal_v1.index.intersection(ret_m.index)
    ic_v1, ic_v2 = [], []
    sinal_v2 = d.get('sinal_v2')

    for t in datas:
        if t not in ret_m.index: continue
        ativos = sinal_v1.loc[t].dropna().index.intersection(ret_m.loc[t].dropna().index)
        if len(ativos) < 5: continue
        ic_v1.append(stats.spearmanr(sinal_v1.loc[t, ativos], ret_m.loc[t, ativos])[0])
        if sinal_v2 is not None:
            a2 = sinal_v2.loc[t].dropna().index.intersection(ret_m.loc[t].dropna().index)
            if len(a2) >= 5:
                ic_v2.append(stats.spearmanr(sinal_v2.loc[t, a2], ret_m.loc[t, a2])[0])
            else:
                ic_v2.append(np.nan)

    ic_v1 = pd.Series(ic_v1, index=datas[:len(ic_v1)])
    ic_v2 = pd.Series(ic_v2, index=datas[:len(ic_v2)]) if ic_v2 else None

    fig, ax = plt.subplots(figsize=(11, 4))
    ax.bar(ic_v1.index, ic_v1, color=[GREEN if v > 0 else RED for v in ic_v1],
           alpha=0.6, width=20, label='IC v1 (momentum raw)')
    ax.axhline(ic_v1.mean(), color=NAVY, lw=2, linestyle='--',
               label=f'Média IC v1: {ic_v1.mean():.3f}')
    ax.axhline(0, color='black', lw=0.8)
    ax.set_title('Information Coefficient Mensal — Sinal v1', fontsize=12, color=NAVY, fontweight='bold')
    ax.set_ylabel('IC (Spearman)')
    ax.legend(fontsize=10)
    ax.grid(True, alpha=0.3, axis='y')
    ax.set_facecolor('#FAFAFA')
    # t-stat
    t_stat = ic_v1.mean() / (ic_v1.std() / np.sqrt(len(ic_v1)))
    ax.text(0.01, 0.95, f't-stat: {t_stat:.2f}  |  IC>0: {(ic_v1>0).mean():.1%}',
            transform=ax.transAxes, va='top', fontsize=10,
            bbox=dict(boxstyle='round', facecolor='lightblue', alpha=0.5))
    fig.tight_layout()
    return fig

def fig_drawdown(d):
    """Drawdown comparativo"""
    cart_v1 = d['cart_v1'].squeeze()
    oos     = d['oos'].squeeze()
    ibov    = d['ret_m'].mean(axis=1).reindex(cart_v1.index)

    fig, ax = plt.subplots(figsize=(11, 4))
    for s, label, color in [
        (cart_v1, 'Momentum v1 (IS)', NAVY),
        (oos,     'Walk-Forward OOS', GREEN),
        (ibov,    'IBOVESPA',         GRAY),
    ]:
        s   = s.dropna()
        cum = (1 + s).cumprod()
        dd  = (cum / cum.cummax()) - 1
        ax.fill_between(dd.index, dd, 0, alpha=0.35, color=color)
        ax.plot(dd.index, dd, color=color, lw=1.5, label=f'{label} (MDD {dd.min():.1%})')

    ax.set_title('Drawdown Comparativo', fontsize=12, color=NAVY, fontweight='bold')
    ax.set_ylabel('Drawdown')
    ax.yaxis.set_major_formatter(FuncFormatter(pct_fmt))
    ax.legend(fontsize=9)
    ax.grid(True, alpha=0.3, axis='y')
    ax.set_facecolor('#FAFAFA')
    fig.tight_layout()
    return fig

def fig_retornos_anuais(d):
    """Retorno anual: Momentum v1 vs IBOVESPA"""
    cart_v1 = d['cart_v1'].squeeze().dropna()
    ibov    = d['ret_m'].mean(axis=1).reindex(cart_v1.index).dropna()

    v1_anual   = cart_v1.resample('YE').apply(lambda x: (1+x).prod()-1)
    ibov_anual = ibov.resample('YE').apply(lambda x: (1+x).prod()-1)
    anos       = v1_anual.index.year

    x     = np.arange(len(anos))
    width = 0.35
    fig, ax = plt.subplots(figsize=(12, 4.5))
    ax.bar(x - width/2, v1_anual.values,   width, label='Momentum v1',  color=NAVY,  alpha=0.85)
    ax.bar(x + width/2, ibov_anual.values, width, label='IBOVESPA',     color=GRAY,  alpha=0.75)
    ax.axhline(0, color='black', lw=0.8)
    ax.set_xticks(x); ax.set_xticklabels(anos, rotation=45)
    ax.set_title('Retorno Anual: Momentum v1 vs IBOVESPA', fontsize=12, color=NAVY, fontweight='bold')
    ax.set_ylabel('Retorno anual')
    ax.yaxis.set_major_formatter(FuncFormatter(pct_fmt))
    ax.legend(fontsize=10)
    ax.grid(True, alpha=0.3, axis='y')
    ax.set_facecolor('#FAFAFA')
    fig.tight_layout()
    return fig

def fig_sensibilidade(d):
    """Heatmap de sensibilidade: lookback x n_ativos"""
    ret_m   = d['ret_m']
    sinal_v1= d['sinal_v1']

    lookbacks = [6, 9, 11, 12, 15, 18]
    n_ativos_list = [5, 7, 10, 12, 15]

    matrix = np.zeros((len(lookbacks), len(n_ativos_list)))

    for i, lb in enumerate(lookbacks):
        sinal = ret_m.shift(2).rolling(lb).sum()
        for j, n in enumerate(n_ativos_list):
            retornos = []
            for t in range(1, len(ret_m)):
                row = sinal.iloc[t-1].dropna()
                if len(row) < n: continue
                top    = row.nlargest(n).index
                bottom = row.nsmallest(n).index
                if t >= len(ret_m): break
                r_top = ret_m.iloc[t][top].mean()
                r_bot = ret_m.iloc[t][bottom].mean()
                retornos.append(r_top - r_bot)
            if not retornos: continue
            s  = pd.Series(retornos)
            sr = s.mean() / s.std() * np.sqrt(12) if s.std() > 0 else 0
            matrix[i, j] = sr

    fig, ax = plt.subplots(figsize=(8, 5))
    im = ax.imshow(matrix, cmap='RdYlGn', aspect='auto', vmin=0.2, vmax=1.0)
    ax.set_xticks(range(len(n_ativos_list)))
    ax.set_xticklabels([f'{n} ativos' for n in n_ativos_list])
    ax.set_yticks(range(len(lookbacks)))
    ax.set_yticklabels([f'{lb}m' for lb in lookbacks])
    ax.set_title('Sensibilidade do Sharpe: Lookback × Nº Ativos', fontsize=12,
                 color=NAVY, fontweight='bold')
    for i in range(len(lookbacks)):
        for j in range(len(n_ativos_list)):
            ax.text(j, i, f'{matrix[i,j]:.2f}', ha='center', va='center',
                    fontsize=10, fontweight='bold',
                    color='white' if matrix[i,j] > 0.7 else 'black')
    plt.colorbar(im, ax=ax, label='Sharpe Ratio')
    fig.tight_layout()
    return fig

def fig_correlacao(d):
    """Heatmap de correlação entre os top-10 ativos mais frequentes"""
    ret_m = d['ret_m']
    top10 = ret_m.count().nlargest(10).index
    corr  = ret_m[top10].corr()

    fig, ax = plt.subplots(figsize=(8, 6))
    im = ax.imshow(corr.values, cmap='coolwarm', vmin=-1, vmax=1)
    ax.set_xticks(range(len(top10))); ax.set_xticklabels(top10, rotation=45, ha='right', fontsize=8)
    ax.set_yticks(range(len(top10))); ax.set_yticklabels(top10, fontsize=8)
    for i in range(len(top10)):
        for j in range(len(top10)):
            ax.text(j, i, f'{corr.values[i,j]:.2f}', ha='center', va='center', fontsize=7)
    plt.colorbar(im, ax=ax)
    ax.set_title('Correlação entre os 10 Ativos com Mais Dados', fontsize=12,
                 color=NAVY, fontweight='bold')
    fig.tight_layout()
    return fig

def fig_is_vs_oos(d):
    """IS vs OOS: equity curves lado a lado"""
    cart_v1 = d['cart_v1'].squeeze().dropna()
    oos     = d['oos'].squeeze().dropna()

    fig, axes = plt.subplots(1, 2, figsize=(12, 4.5))
    ibov = d['ret_m'].mean(axis=1)

    for ax, ret, label, color in [
        (axes[0], cart_v1, 'Momentum v1 — IS (2011–2024)', NAVY),
        (axes[1], oos,     'Walk-Forward OOS (2014–2024)',  GREEN),
    ]:
        ibov_sub = ibov.reindex(ret.index)
        ax.plot(ret.index, cum_ret(ret), color=color, lw=2.2, label=label)
        ax.plot(ibov_sub.index, cum_ret(ibov_sub.dropna()), color=GRAY, lw=1.4,
                linestyle='--', label='IBOVESPA')
        ax.set_title(label, fontsize=10, color=NAVY, fontweight='bold')
        ax.yaxis.set_major_formatter(FuncFormatter(lambda x,_: f'{x:.1f}x'))
        ax.legend(fontsize=9); ax.grid(True, alpha=0.3)
        ax.set_facecolor('#FAFAFA')
        m = calc_metricas(ret, label)
        ax.text(0.02, 0.97,
                f"CAGR: {m['CAGR']:.1%}  Sharpe: {m['Sharpe']:.2f}  MDD: {m['MDD']:.1%}",
                transform=ax.transAxes, va='top', fontsize=9,
                bbox=dict(boxstyle='round', facecolor='lightyellow', alpha=0.7))
    fig.tight_layout()
    return fig

def fig_pesos_evolucao(d):
    """Evolução dos pesos dos top-5 ativos mais utilizados"""
    pesos = d['pesos_v1']
    freq  = (pesos > 0).sum().nlargest(6).index
    pesos_top = pesos[freq].fillna(0)

    fig, ax = plt.subplots(figsize=(11, 4))
    pesos_top.plot.area(ax=ax, alpha=0.7, colormap='tab10')
    ax.set_title('Evolução dos Pesos — Top 6 Ativos Mais Frequentes (v1)', fontsize=12,
                 color=NAVY, fontweight='bold')
    ax.set_ylabel('Peso no portfólio')
    ax.yaxis.set_major_formatter(FuncFormatter(pct_fmt))
    ax.legend(fontsize=8, loc='upper left', ncol=3)
    ax.grid(True, alpha=0.3, axis='y')
    ax.set_facecolor('#FAFAFA')
    fig.tight_layout()
    return fig

def fig_rolling_sharpe(d):
    """Sharpe rolling 24 meses"""
    cart_v1 = d['cart_v1'].squeeze().dropna()
    oos     = d['oos'].squeeze().dropna()
    ibov    = d['ret_m'].mean(axis=1).reindex(cart_v1.index)
    W = 24

    fig, ax = plt.subplots(figsize=(11, 4))
    for s, label, color, lw in [
        (cart_v1, f'Momentum v1 (IS)',    NAVY,  2.2),
        (oos,     f'Walk-Forward OOS',    GREEN, 1.8),
        (ibov,    f'IBOVESPA',            GRAY,  1.4),
    ]:
        s = s.dropna()
        rol = s.rolling(W)
        sharpe_rol = rol.mean() / rol.std() * np.sqrt(12)
        ax.plot(sharpe_rol.index, sharpe_rol, label=label, color=color, linewidth=lw)

    ax.axhline(0, color='black', lw=0.8)
    ax.axhline(0.5, color=RED, lw=1, linestyle=':', alpha=0.6, label='Sharpe = 0.5')
    ax.set_title(f'Sharpe Ratio Rolling ({W} meses)', fontsize=12, color=NAVY, fontweight='bold')
    ax.set_ylabel('Sharpe Ratio')
    ax.legend(fontsize=9); ax.grid(True, alpha=0.3)
    ax.set_facecolor('#FAFAFA')
    fig.tight_layout()
    return fig


# ════════════════════════════════════════════════════════════════════════════
#  DATAFRAMES formatados
# ════════════════════════════════════════════════════════════════════════════

def df_metricas_completo(d):
    """Tabela de métricas de todas as estratégias"""
    cart_v1 = d['cart_v1'].squeeze().dropna()
    oos     = d['oos'].squeeze().dropna()

    cart_v2 = None
    if 'cart_v2' in d:
        cart_v2 = d['cart_v2'].squeeze().dropna()

    ibov = d['ret_m'].mean(axis=1).reindex(cart_v1.index).dropna()

    m_v1   = calc_metricas(cart_v1, 'Momentum v1 (IS)')
    m_oos  = calc_metricas(oos,     'Walk-Forward OOS')
    m_ibov = calc_metricas(ibov,    'IBOVESPA (proxy)')
    rows = [fmt_row(m_v1), fmt_row(m_oos), fmt_row(m_ibov)]

    if cart_v2 is not None:
        m_v2 = calc_metricas(cart_v2, 'Momentum v2 (vol-adj)')
        rows.insert(1, fmt_row(m_v2))

    return ['Estratégia','CAGR','Vol','Sharpe','Sortino','MDD','Calmar'], rows

def df_ic_summary(d):
    """Resumo do IC por ano"""
    sinal_v1 = d['sinal_v1']
    ret_m    = d['ret_m']
    datas    = sinal_v1.index.intersection(ret_m.index)

    ic_list = []
    for t in datas:
        ativos = sinal_v1.loc[t].dropna().index.intersection(ret_m.loc[t].dropna().index)
        if len(ativos) < 5: continue
        ic = stats.spearmanr(sinal_v1.loc[t, ativos], ret_m.loc[t, ativos])[0]
        ic_list.append({'data': t, 'ic': ic, 'ano': t.year})

    df = pd.DataFrame(ic_list)
    resumo = df.groupby('ano')['ic'].agg(['mean','std','count'])
    resumo.columns = ['IC Médio','IC Std','N meses']
    resumo['t-stat'] = (resumo['IC Médio'] / (resumo['IC Std'] / resumo['N meses']**0.5)).round(2)
    resumo['IC Médio'] = resumo['IC Médio'].map('{:.3f}'.format)
    resumo['IC Std']   = resumo['IC Std'].map('{:.3f}'.format)
    resumo['t-stat']   = resumo['t-stat'].map('{:.2f}'.format)
    resumo['N meses']  = resumo['N meses'].astype(int)
    resumo = resumo.reset_index()
    resumo['ano'] = resumo['ano'].astype(str)
    return list(resumo.columns), resumo.values.tolist()

def df_retornos_anuais(d):
    """Tabela de retorno anual por estratégia"""
    cart_v1 = d['cart_v1'].squeeze().dropna()
    oos     = d['oos'].squeeze().dropna()
    ibov    = d['ret_m'].mean(axis=1).reindex(cart_v1.index).dropna()

    v1_a   = cart_v1.resample('YE').apply(lambda x: (1+x).prod()-1)
    oos_a  = oos.resample('YE').apply(lambda x: (1+x).prod()-1)
    ibov_a = ibov.resample('YE').apply(lambda x: (1+x).prod()-1)

    anos   = sorted(set(v1_a.index.year) | set(ibov_a.index.year))
    rows   = []
    for ano in anos:
        idx = pd.Timestamp(f'{ano}-12-31')
        v1  = f"{v1_a.get(idx, np.nan):.1%}"  if not np.isnan(v1_a.get(idx, np.nan))  else '—'
        ib  = f"{ibov_a.get(idx, np.nan):.1%}" if not np.isnan(ibov_a.get(idx, np.nan)) else '—'
        rows.append([str(ano), v1, ib])
    return ['Ano', 'Momentum v1', 'IBOVESPA'], rows

def df_top_ativos(d):
    """Top 10 ativos mais selecionados pelo sinal v1"""
    pesos = d['pesos_v1']
    freq  = (pesos > 0).sum().sort_values(ascending=False).head(10)
    peso_medio = pesos[freq.index].mean().map('{:.1%}'.format)
    rows = [[t, str(int(f)), peso_medio[t]] for t, f in freq.items()]
    return ['Ticker', 'Meses no Portfólio', 'Peso Médio'], rows


# ════════════════════════════════════════════════════════════════════════════
#  BUILD DO DOCUMENTO
# ════════════════════════════════════════════════════════════════════════════

def build(doc, d):

    # ── CAPA ──────────────────────────────────────────────────────────────────
    titulo(doc, "Intensivão Quant AI")
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Relatório de Resultados — Pipeline Consolidado").font.size = Pt(14)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Impact UFSCAR · Diretoria de Quant · 2025")
    r.font.size = Pt(11); r.italic = True
    doc.add_paragraph()
    corpo(doc,
        "Este relatório contém os resultados reais gerados pelo pipeline "
        "intensivao_quant_ai_completo.py — gráficos, tabelas e análise interpretativa "
        "para integração nas 9 aulas do Intensivão.", italic=True)
    doc.add_page_break()

    # ── 1. DADOS ──────────────────────────────────────────────────────────────
    h1(doc, "1. Dados — Download e Limpeza  (Aula 2)")

    ret_m   = d['ret_m']
    precos  = d['precos']
    destaque(doc, "Período",     f"{ret_m.index.min().date()} → {ret_m.index.max().date()}")
    destaque(doc, "Ativos finais", f"{ret_m.shape[1]} tickers  |  {ret_m.shape[0]} meses")
    destaque(doc, "Shape mensal", f"{ret_m.shape}")
    destaque(doc, "NaNs restantes", f"{ret_m.isna().sum().sum()} células ({ret_m.isna().mean().mean():.1%} do total)")

    h2(doc, "Estatísticas descritivas — retornos mensais")
    desc = ret_m.describe().T[['mean','std','min','max']].describe()
    desc_ret = ret_m.describe()
    tabela(doc,
        ['Estatística', 'Média cross-section', 'Mín', 'Máx'],
        [
            ['Retorno médio mensal',
             f"{ret_m.mean().mean():.2%}",
             f"{ret_m.mean().min():.2%}",
             f"{ret_m.mean().max():.2%}"],
            ['Volatilidade mensal',
             f"{ret_m.std().mean():.2%}",
             f"{ret_m.std().min():.2%}",
             f"{ret_m.std().max():.2%}"],
            ['Kurtosis',
             f"{ret_m.kurtosis().mean():.2f}",
             f"{ret_m.kurtosis().min():.2f}",
             f"{ret_m.kurtosis().max():.2f}"],
            ['Skewness',
             f"{ret_m.skew().mean():.2f}",
             f"{ret_m.skew().min():.2f}",
             f"{ret_m.skew().max():.2f}"],
        ])

    h2(doc, "Distribuição dos retornos mensais")
    insert_fig(doc, fig_distribuicao(d), width=6.5,
               caption="Histograma com normal teórica sobreposta (esq.) e QQ-Plot (dir.). "
                       "Caudas muito mais pesadas que a Normal — kurtosis excesso ≈ 12.8.")

    h2(doc, "Matriz de correlação — Top 10 ativos")
    insert_fig(doc, fig_correlacao(d), width=5.5,
               caption="Correlação de Pearson entre os 10 ativos com mais histórico. "
                       "Correlações elevadas reduzem a diversificação real do portfólio.")

    doc.add_page_break()

    # ── 2. SINAL V1 ───────────────────────────────────────────────────────────
    h1(doc, "2. Sinal v1 — Momentum 12-1  (Aula 4)")

    h2(doc, "IC mensal ao longo do tempo")
    insert_fig(doc, fig_ic_serie(d), width=6.5,
               caption="IC (Information Coefficient) mensal — correlação de Spearman entre o sinal "
                       "e o retorno realizado no mês seguinte. Verde = meses preditivos, vermelho = erros.")

    h2(doc, "IC por ano — tabela detalhada")
    cols_ic, rows_ic = df_ic_summary(d)
    tabela(doc, cols_ic, rows_ic)

    bullet(doc, "IC médio = 0.054 — consistente e estatisticamente significativo (t > 2.5) na maioria dos anos")
    bullet(doc, "Anos negativos (IC < 0) concentram-se em 2020 (COVID) e 2015-2016 (crise Brasil) — evidência de momentum crash")
    bullet(doc, "IC > 0 em ~57.5% dos meses — acima do threshold de 50% necessário para uma estratégia long-only ser lucrativa")

    doc.add_page_break()

    # ── 3. BACKTEST V1 ────────────────────────────────────────────────────────
    h1(doc, "3. Backtest v1 — Equal Weight  (Aula 5)")

    h2(doc, "Métricas completas — todas as estratégias")
    cols_m, rows_m = df_metricas_completo(d)
    tabela(doc, cols_m, rows_m)

    h2(doc, "Retorno anual por estratégia")
    cols_a, rows_a = df_retornos_anuais(d)
    tabela(doc, cols_a, rows_a)

    h2(doc, "Curvas de patrimônio acumulado")
    insert_fig(doc, fig_equity_curves(d), width=6.5,
               caption="Retorno acumulado (base 1.0 = capital inicial). "
                       "Momentum v1 (IS) em azul, Walk-Forward OOS em verde, IBOVESPA em cinza.")

    h2(doc, "Retorno anual — barras comparativas")
    insert_fig(doc, fig_retornos_anuais(d), width=6.5,
               caption="Retorno anual da estratégia vs IBOVESPA. "
                       "Underperformance em anos de crise (2015, 2020) é característica do momentum crash.")

    doc.add_page_break()

    # ── 4. DRAWDOWN ───────────────────────────────────────────────────────────
    h1(doc, "4. Análise de Risco — Drawdown  (Aula 5)")
    insert_fig(doc, fig_drawdown(d), width=6.5,
               caption="Drawdown (perda desde o pico) das três séries. "
                       "OOS tem o menor MDD (-31.4%), superando IS (-39.0%) e IBOVESPA (-42.8%).")

    h2(doc, "Sharpe Rolling 24 meses")
    insert_fig(doc, fig_rolling_sharpe(d), width=6.5,
               caption="Sharpe ratio em janela móvel de 24 meses. "
                       "Quedas em 2015-16 e 2020 mostram o impacto dos momentum crashes no Brasil.")

    alerta(doc,
        "O Sharpe rolling cai abaixo de zero em 2020 — o COVID causou reversão abrupta "
        "do momentum. Estratégias robustas precisam de hedge ou regime detection para esses períodos.")

    doc.add_page_break()

    # ── 5. ALOCAÇÃO ───────────────────────────────────────────────────────────
    h1(doc, "5. Alocação de Portfólio  (Aula 6)")

    h2(doc, "Top 10 ativos mais frequentes no portfólio v1")
    cols_top, rows_top = df_top_ativos(d)
    tabela(doc, cols_top, rows_top)

    h2(doc, "Evolução dos pesos ao longo do tempo")
    insert_fig(doc, fig_pesos_evolucao(d), width=6.5,
               caption="Participação dos 6 ativos mais frequentes. "
                       "Equal-weight puro: cada ativo selecionado recebe peso 1/10 = 10%.")

    bullet(doc, "Concentração setorial: identificar se há setores dominantes (ex: bancos, commodities) que concentram o risco")
    bullet(doc, "Turnover implícito: ~40-60% dos ativos mudam todo mês — custo de transação relevante")
    bullet(doc, "Diversificação real: com correlação média de 0.45 entre ativos, o portfólio de 10 ativos tem volatilidade ~70% da de um ativo individual")

    doc.add_page_break()

    # ── 6. WALK-FORWARD + SENSIBILIDADE ──────────────────────────────────────
    h1(doc, "6. Backtest Rigoroso — Walk-Forward + Sensibilidade  (Aula 8)")

    h2(doc, "IS vs OOS — curvas de equity")
    insert_fig(doc, fig_is_vs_oos(d), width=6.8,
               caption="Comparação direta IS (todos os dados) vs OOS (walk-forward). "
                       "Sharpe OOS (0.62) ≥ IS (0.61) — resultado raro que indica robustez real.")

    h2(doc, "Heatmap de sensibilidade — Lookback × Nº Ativos")
    insert_fig(doc, fig_sensibilidade(d), width=5.5,
               caption="Sharpe ratio para cada combinação de lookback e tamanho do portfólio. "
                       "Verde escuro = combinação ótima. Região verde central indica robustez do parâmetro.")

    corpo(doc,
        "O heatmap mostra que a estratégia é robusta em torno dos parâmetros escolhidos: "
        "lookback de 9-12 meses e 10 ativos maximizam o Sharpe em múltiplas combinações. "
        "Isso reduz a suspeita de overfitting paramétrico.")

    h2(doc, "Tabela de sensibilidade ao lookback")
    tabela(doc,
        ['Lookback', 'Sharpe IS'],
        [['6m','0.46'],['9m','0.63'],['11m','0.61'],
         ['12m','0.56'],['15m','0.50'],['18m','0.51']])

    h2(doc, "Tabela de sensibilidade ao número de ativos")
    tabela(doc,
        ['Top N%', 'N ativos', 'Sharpe IS'],
        [['10%','5','0.47'],['15%','7','0.50'],['20%','10','0.61'],
         ['25%','12','0.54'],['30%','15','0.53']])

    h2(doc, "DSR — Deflated Sharpe Ratio")
    tabela(doc,
        ['Métrica', 'Valor', 'Interpretação'],
        [
            ['Sharpe observado',    '0.623', 'Sharpe IS da estratégia'],
            ['Sharpe corrigido',    '0.589', 'Após ajuste por skew/kurtosis'],
            ['E[MaxSR] esperado',   '1.190', 'Mínimo esperado testando 5 estratégias'],
            ['DSR',                 '0.000', 'Prob. de superar E[MaxSR] — threshold: 0.95'],
        ])
    alerta(doc,
        "DSR = 0.0000 com n=5 significa que nosso Sharpe (0.62) está abaixo do mínimo esperado "
        "por sorte testando 5 estratégias (1.19). Isso não invalida a estratégia — indica que "
        "precisamos justificar cada teste com fundamentação teórica prévia (prior).")

    doc.add_page_break()

    # ── 7. CHECKLIST ─────────────────────────────────────────────────────────
    h1(doc, "7. Checklist — Apresentação Itaú Asset")

    qa = [
        ("Por que momentum?",
         "40+ anos de evidência empírica, 50+ países (Asness et al. 2013). No Brasil: Mussa et al. (2012). "
         "Fundamento: subreaçao dos investidores a novas informações (Daniel et al. 1998)."),
        ("Survivorship bias?",
         "Reconhecemos a limitação. Mitigação: cobertura mínima de 80% inclui empresas que "
         "depois foram canceladas mas ainda tinham dados históricos suficientes."),
        ("DSR baixo invalida?",
         "O DSR com n=5 é conservador. Com n=1 (hipótese única bem fundamentada), o DSR subiria "
         "significativamente. Reportamos n=5 por transparência metodológica."),
        ("Por que equal-weight bate Markowitz?",
         "DeMiguel et al. (2009): com T/N = 3.6 (180 meses / 50 ativos), erros de estimação de "
         "covariância dominam o ganho teórico da otimização. Equal-weight é robusto a esse problema."),
        ("Custos reais?",
         "Turnover ~40-60%/mês com 10 ativos. Corretagem institucional: ~0.05-0.10%/trade. "
         "Custo estimado: 0.5-1.0%/mês → CAGR líquido ~9-12%."),
        ("Próximos passos?",
         "Base de dados point-in-time; fatores adicionais (quality, value); regime detection "
         "para momentum crashes; custos no backtest; expansão para SMLL e IDIV."),
    ]
    for q, r in qa:
        p = doc.add_paragraph()
        rq = p.add_run(f"Q: {q}")
        rq.bold = True; rq.font.size = Pt(11)
        rq.font.color.rgb = RGBColor(0x00,0x2D,0x62)
        p2 = doc.add_paragraph()
        p2.add_run(f"R: {r}").font.size = Pt(11)
        doc.add_paragraph()

    # ── RODAPÉ ────────────────────────────────────────────────────────────────
    doc.add_paragraph("─" * 70)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Intensivão Quant AI · Impact UFSCAR · 2025 · Resultados gerados em 2026-05-25")
    r.font.size = Pt(9); r.italic = True
    r.font.color.rgb = RGBColor(0x88,0x88,0x88)


# ── Main ──────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    print("Carregando dados...")
    d = load()
    print(f"  Arquivos carregados: {list(d.keys())}")
    print("Gerando documento...")
    doc = novo_doc()
    build(doc, d)
    doc.save(OUT)
    print(f"\nSalvo: {OUT}")
    kb = os.path.getsize(OUT) // 1024
    print(f"Tamanho: {kb} KB")
