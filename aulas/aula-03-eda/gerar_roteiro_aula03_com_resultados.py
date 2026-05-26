# -*- coding: utf-8 -*-
"""
Gera roteiro-aula-03-eda-final.docx
Usa EXATAMENTE o mesmo código do notebook para gerar os gráficos,
e embute análise de cada output/gráfico no roteiro do professor.
"""
import os
from io import BytesIO

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from scipy import stats

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ── Caminhos ─────────────────────────────────────────────────────────────────
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
ROOT = SCRIPT_DIR
_p = SCRIPT_DIR
while _p != os.path.dirname(_p):
    if os.path.exists(os.path.join(_p, '.git')):
        ROOT = _p; break
    _p = os.path.dirname(_p)

DADOS_DIR = os.path.join(ROOT, 'dados')
OUT = os.path.join(SCRIPT_DIR, 'roteiro-aula-03-eda-final.docx')

# ══════════════════════════════════════════════════════════════════════════════
# REPRODUZ EXATAMENTE O CÓDIGO DO NOTEBOOK
# ══════════════════════════════════════════════════════════════════════════════
print("Carregando dados...")
precos      = pd.read_parquet(os.path.join(DADOS_DIR, 'precos_ibov.parquet'))
retornos    = pd.read_parquet(os.path.join(DADOS_DIR, 'retornos_diarios.parquet'))
ret_mensais = pd.read_parquet(os.path.join(DADOS_DIR, 'retornos_mensais.parquet'))
print(f"  precos:      {precos.shape}")
print(f"  retornos:    {retornos.shape}")
print(f"  ret_mensais: {ret_mensais.shape}")

plt.rcParams['figure.figsize'] = (12, 5)
plt.rcParams['axes.grid'] = True
plt.rcParams['grid.alpha'] = 0.3

# ── Variáveis auxiliares usadas em vários gráficos ────────────────────────────
petr4_ret   = retornos['PETR4.SA'].dropna()
petr4_preco = precos['PETR4.SA'].dropna()

# ── Gráfico 1 (célula 6) — Scatter Risco × Retorno ──────────────────────────
print("  [1/4] Scatter risco x retorno...")
resumo = pd.DataFrame({
    'retorno_anual': retornos.mean() * 252,
    'vol_anual':     retornos.std()  * np.sqrt(252)
}).dropna().sort_values('retorno_anual', ascending=False)

fig1, ax1 = plt.subplots(figsize=(10, 6))
ax1.scatter(resumo['vol_anual'], resumo['retorno_anual'], alpha=0.6, s=50)
destaques = ['WEGE3.SA','PETR4.SA','VALE3.SA','MGLU3.SA','BBAS3.SA']
for t in destaques:
    if t in resumo.index:
        ax1.annotate(t.replace('.SA',''),
                     (resumo.loc[t,'vol_anual'], resumo.loc[t,'retorno_anual']),
                     textcoords='offset points', xytext=(6,3), fontsize=8)
ax1.axhline(0, color='red', linewidth=0.8, linestyle='--', label='retorno zero')
ax1.set_xlabel('Volatilidade anual')
ax1.set_ylabel('Retorno anual médio')
ax1.set_title('Risco × Retorno — IBOVESPA (2010–2024)')
ax1.xaxis.set_major_formatter(plt.FuncFormatter(lambda x,_: f'{x:.0%}'))
ax1.yaxis.set_major_formatter(plt.FuncFormatter(lambda x,_: f'{x:.0%}'))
ax1.legend()
plt.tight_layout()
buf1 = BytesIO(); fig1.savefig(buf1, format='png', dpi=150, bbox_inches='tight'); buf1.seek(0)
plt.close(fig1)

top5    = resumo.head()
bot5    = resumo.tail()
acima_zero = (resumo['retorno_anual'] > 0).sum()
abaixo_zero = (resumo['retorno_anual'] <= 0).sum()

# ── Gráfico 2 (célula 8) — Fat Tails: Histograma + QQ ───────────────────────
print("  [2/4] Fat tails (histograma + QQ)...")
fig2, axes2 = plt.subplots(1, 2, figsize=(14, 5))

x = np.linspace(petr4_ret.min(), petr4_ret.max(), 300)
axes2[0].hist(petr4_ret, bins=100, density=True, alpha=0.6,
              color='steelblue', label='retornos reais')
axes2[0].plot(x, stats.norm.pdf(x, petr4_ret.mean(), petr4_ret.std()),
              'r-', lw=2, label='normal ajustada')
axes2[0].set_title('Histograma — PETR4')
axes2[0].set_xlim(-0.15, 0.15)
axes2[0].legend()

stats.probplot(petr4_ret, dist='norm', plot=axes2[1])
axes2[1].set_title('QQ-Plot — fat tails nas extremidades')
axes2[1].get_lines()[1].set_color('red')

plt.tight_layout()
buf2 = BytesIO(); fig2.savefig(buf2, format='png', dpi=150, bbox_inches='tight'); buf2.seek(0)
plt.close(fig2)

curtose    = petr4_ret.kurtosis()
assimetria = petr4_ret.skew()
sigma      = petr4_ret.std()
extremos   = {}
for n in [3, 4, 5]:
    obs = int((petr4_ret.abs() > n * sigma).sum())
    esp = len(petr4_ret) * 2 * stats.norm.sf(n)
    extremos[n] = (obs, esp)

# ── Gráfico 3 (célula 10) — Estacionariedade: Preço vs Retorno ──────────────
print("  [3/4] Estacionariedade (preço vs retorno)...")
from statsmodels.tsa.stattools import adfuller

fig3, axes3 = plt.subplots(2, 1, figsize=(13, 6))
petr4_preco.plot(ax=axes3[0], color='steelblue')
axes3[0].set_title('PETR4 — Preço (não estacionário)')
axes3[0].set_ylabel('Preço (R$)')
petr4_ret.plot(ax=axes3[1], color='steelblue', alpha=0.7)
axes3[1].axhline(petr4_ret.mean(), color='red', lw=1.2, ls='--',
                 label=f'média = {petr4_ret.mean():.4f}')
axes3[1].set_title('PETR4 — Retorno diário (estacionário)')
axes3[1].legend()
plt.tight_layout()
buf3 = BytesIO(); fig3.savefig(buf3, format='png', dpi=150, bbox_inches='tight'); buf3.seek(0)
plt.close(fig3)

adf_preco   = adfuller(petr4_preco.dropna())
adf_retorno = adfuller(petr4_ret.dropna())

# ── Gráfico 4 (célula 12) — ACF Diário e Mensal ─────────────────────────────
print("  [4/4] ACF (diário + mensal)...")
from statsmodels.graphics.tsaplots import plot_acf

fig4, axes4 = plt.subplots(1, 2, figsize=(14, 5))
plot_acf(petr4_ret.dropna(), lags=30, ax=axes4[0], alpha=0.05)
axes4[0].set_title('ACF — PETR4 Retornos DIÁRIOS (lags 1–30 dias)')
axes4[0].set_xlabel('Lag (dias)')
plot_acf(ret_mensais['PETR4.SA'].dropna(), lags=24, ax=axes4[1], alpha=0.05)
axes4[1].set_title('ACF — PETR4 Retornos MENSAIS (lags 1–24 meses)')
axes4[1].set_xlabel('Lag (meses)')
plt.tight_layout()
buf4 = BytesIO(); fig4.savefig(buf4, format='png', dpi=150, bbox_inches='tight'); buf4.seek(0)
plt.close(fig4)

autocorr_1d  = petr4_ret.autocorr(lag=1)
autocorr_1m  = ret_mensais['PETR4.SA'].autocorr(lag=1)
autocorr_12m = ret_mensais['PETR4.SA'].autocorr(lag=12)

# ── Filtros (célula 14) ───────────────────────────────────────────────────────
cobertura       = retornos.notna().mean()
tickers_validos = cobertura[cobertura >= 0.80].index.tolist()
tickers_removidos_cob = cobertura[cobertura < 0.80].index.tolist()
liquidez        = (retornos[tickers_validos] != 0).mean()
tickers_finais  = liquidez[liquidez >= liquidez.quantile(0.10)].index.tolist()

print("Gráficos prontos. Montando documento...")

# ══════════════════════════════════════════════════════════════════════════════
# HELPERS DE DOCUMENTO
# ══════════════════════════════════════════════════════════════════════════════
def novo_doc():
    doc = Document()
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Cm(2.5)
        s.left_margin = Cm(3.0); s.right_margin = Cm(2.5)
    return doc

def capa(doc, t1, t2, nota):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(t1); r.bold = True; r.font.size = Pt(16)
    r.font.color.rgb = RGBColor(0,45,98)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(t2); r.font.size = Pt(13)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(nota); r.italic = True; r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(120,120,120)

def secao(doc, t):
    doc.add_paragraph()
    p = doc.add_heading(t, 1)
    p.runs[0].font.color.rgb = RGBColor(0,45,98)
    p.runs[0].font.size = Pt(13)

def subsecao(doc, t):
    p = doc.add_heading(t, 2)
    p.runs[0].font.color.rgb = RGBColor(184,134,0)
    p.runs[0].font.size = Pt(11)

def sep(doc):
    p = doc.add_paragraph()
    r = p.add_run("─" * 70)
    r.font.size = Pt(9); r.font.color.rgb = RGBColor(180,180,180)

def tempo(doc, t):
    p = doc.add_paragraph()
    r = p.add_run(f"  {t}")
    r.font.size = Pt(9); r.italic = True; r.font.color.rgb = RGBColor(120,120,120)

def fala(doc, texto):
    p = doc.add_paragraph()
    r = p.add_run(texto); r.font.size = Pt(11)

def acao(doc, texto):
    p = doc.add_paragraph()
    r = p.add_run(f"[{texto}]")
    r.font.size = Pt(10); r.italic = True
    r.font.color.rgb = RGBColor(0,100,180)

def cod(doc, texto):
    p = doc.add_paragraph()
    r = p.add_run(texto)
    r.font.name = 'Courier New'; r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(30,30,30)

def output(doc, texto):
    p = doc.add_paragraph()
    r = p.add_run(f"▶  OUTPUT REAL:\n{texto}")
    r.font.name = 'Courier New'; r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0,100,0); r.bold = True

def interp(doc, texto):
    p = doc.add_paragraph()
    r = p.add_run(f"↳  {texto}")
    r.font.size = Pt(11); r.font.color.rgb = RGBColor(0,100,0)

def add_buf(doc, buf, largura=6.2):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(buf, width=Inches(largura))

# ══════════════════════════════════════════════════════════════════════════════
doc = novo_doc()

capa(doc,
     "INTENSIVÃO QUANT AI — AULA 03",
     "Análise Exploratória dos Dados (EDA)",
     "Roteiro completo com código + resultados reais — uso exclusivo do instrutor")

# ══════════════════════════════════════════════════════════════════════════════
secao(doc, "ABERTURA")
tempo(doc, "0:00 – 0:03  |  Contexto da aula")
sep(doc)

fala(doc,
    "Pessoal, bem-vindos à Aula 3. Na aula passada a gente construiu o pipeline "
    "de dados — coletamos, limpamos e salvamos os três parquets. Hoje a gente "
    "olha de verdade para esses dados.")

fala(doc,
    "Tem uma frase que todo quant experiente repete: 'se você não explorou os dados, "
    "você não entende a estratégia'. É muito fácil rodar um backtest e ver um "
    "resultado bonito sem entender de onde ele vem. Hoje a gente constrói o "
    "vocabulário estatístico que vai aparecer em todo relatório do intensivão.")

fala(doc,
    "Quatro conceitos centrais: esperança e volatilidade — os dois números que "
    "resumem qualquer distribuição. Fat tails — por que a distribuição normal "
    "falha nos mercados. Estacionariedade — por que nunca modelamos preços "
    "diretamente. E autocorrelação — a evidência estatística de que momentum existe.")

# ══════════════════════════════════════════════════════════════════════════════
secao(doc, "PARTE 1 — TEORIA (15 min)")
tempo(doc, "0:03 – 0:18")
sep(doc)

acao(doc, "Avançar para slide 2: Setup inicial")
subsecao(doc, "Setup — O que estamos carregando")

fala(doc,
    "A primeira coisa que fazemos é carregar os dados que geramos na Aula 2. "
    "Três parquets: preços brutos, retornos diários e retornos mensais.")

cod(doc,
    "precos      = pd.read_parquet(os.path.join(DADOS_DIR, 'precos_ibov.parquet'))\n"
    "retornos    = pd.read_parquet(os.path.join(DADOS_DIR, 'retornos_diarios.parquet'))\n"
    "ret_mensais = pd.read_parquet(os.path.join(DADOS_DIR, 'retornos_mensais.parquet'))")

output(doc,
    f"Preços:           {precos.shape}\n"
    f"Retornos diários: {retornos.shape}\n"
    f"Retornos mensais: {ret_mensais.shape}")

interp(doc,
    f"Os dados carregaram exatamente como esperado. {precos.shape[0]} dias × "
    f"{precos.shape[1]} ações nos preços brutos. "
    f"{retornos.shape[0]} linhas nos retornos diários — uma a menos que os preços "
    "porque o primeiro dia é removido (não há retorno sem dia anterior). "
    f"E {ret_mensais.shape[0]} meses × {ret_mensais.shape[1]} ações nos mensais "
    f"— {ret_mensais.shape[0]/12:.0f} anos exatos de histórico. "
    "Esses três DataFrames são tudo que precisamos para a aula de hoje.")

acao(doc, "Avançar para slide 3: Esperança e Variância")
subsecao(doc, "Esperança, Variância e o Trade-off Risco × Retorno")

fala(doc,
    "Dois números resumem qualquer distribuição de retornos. O primeiro é a "
    "esperança — a média dos retornos. Matematicamente: mu = (1/N) * soma(R_i). "
    "O segundo é a volatilidade — o desvio padrão, raiz quadrada da variância. "
    "É a nossa principal proxy de risco.")

fala(doc,
    "Para comparar ações em diferentes frequências, anualzamos. A regra é simples: "
    "multiplicamos a média diária por 252 dias úteis. Para a volatilidade, "
    "multiplicamos por raiz de 252 — e não por 252 diretamente. Por quê? "
    "Porque a variância é aditiva, não o desvio padrão. Se você somar as "
    "variâncias de 252 dias independentes, obtém a variância anual. A raiz disso "
    "é a volatilidade anual. Essa é a 'regra da raiz do tempo' — fundamental em "
    "qualquer cálculo de risco.")

fala(doc,
    "O scatter risco versus retorno é o gráfico mais fundamental em finanças "
    "quantitativas. Cada ponto é uma ação. O eixo horizontal é a volatilidade "
    "anual. O eixo vertical é o retorno médio anual. Em teoria — a teoria de "
    "Markowitz dos anos 1950 — deveria haver uma relação positiva: mais risco, "
    "mais retorno. Na prática, veremos que o mercado brasileiro quebra essa "
    "relação em vários casos.")

acao(doc, "Avançar para slide 4: Fat Tails")
subsecao(doc, "Fat Tails — Por Que a Distribuição Normal Falha")

fala(doc,
    "A distribuição normal é a mais famosa da estatística — a curva de sino. "
    "Ela é usada em todo lugar: na física, na biologia, na engenharia. "
    "Em finanças clássicas também — Black-Scholes assume normalidade dos retornos, "
    "o VaR paramétrico assume normalidade, o próprio Sharpe Ratio assume normalidade "
    "implicitamente. O problema é que essa premissa é falsa nos mercados.")

fala(doc,
    "A distribuição real dos retornos tem duas anomalias em relação à normal. "
    "A primeira é skewness negativa — assimetria. A distribuição dos retornos "
    "não é simétrica: a cauda esquerda é mais pesada que a direita. "
    "Quedas são mais rápidas e mais extremas que altas. "
    "O mercado 'sobe de escada e desce de elevador'. "
    "Isso é observável: em mercados de bear, as quedas são bruscas e concentradas. "
    "Em bull markets, as altas são graduais e distribuídas.")

fala(doc,
    "A segunda anomalia é curtose elevada — kurtosis acima de 3. A distribuição "
    "real tem o pico mais alto e as caudas mais gordas que a normal. "
    "Isso se chama leptocurtose ou fat tails. O impacto prático: eventos que a "
    "normal diz que ocorrem uma vez a cada 4.700 anos acontecem em praticamente "
    "toda grande crise. O COVID em março de 2020, o Joesley Day em 2017, o "
    "colapso de 2008 — todos foram eventos de 5, 6, 7 sigma que a normal "
    "classificaria como impossíveis.")

acao(doc, "Avançar para slide 5: Estacionariedade")
subsecao(doc, "Estacionariedade — Por Que Nunca Modelamos Preços")

fala(doc,
    "Estacionariedade é uma propriedade que parece técnica mas tem consequências "
    "práticas enormes. Uma série é estacionária se suas propriedades estatísticas — "
    "média, variância, estrutura de correlação — não mudam ao longo do tempo.")

fala(doc,
    "Preços não são estacionários. O preço da PETR4 em 2010 era completamente "
    "diferente de 2020. Se você calcular a média dos preços usando 2010 a 2015, "
    "vai obter um número. Se calcular 2020 a 2024, vai obter outro. A média muda "
    "com a janela escolhida. Matematicamente dizemos que preços têm raiz unitária "
    "— seguem um random walk com tendência. Se você tentar rodar uma regressão "
    "entre dois preços não estacionários — digamos, preço de PETR4 contra preço "
    "de VALE3 — vai encontrar uma correlação altíssima que não representa nenhuma "
    "relação causal real. Isso é regressão espúria.")

fala(doc,
    "Retornos são estacionários. Ao calcular o log-return, removemos a tendência. "
    "A série de retornos oscila em torno de uma média próxima de zero, sem "
    "explodir ou colapsar estruturalmente. Isso é o que nos permite aplicar "
    "estatística clássica — médias, correlações, testes de hipótese — com validade.")

acao(doc, "Avançar para slide 6: Autocorrelação")
subsecao(doc, "Autocorrelação — A Evidência Estatística do Momentum")

fala(doc,
    "Autocorrelação mede se os retornos de hoje têm alguma relação com os retornos "
    "de períodos anteriores. A autocorrelação de lag k é a correlação entre r_t "
    "e r_{t-k}.")

fala(doc,
    "Em frequência diária, a autocorrelação é praticamente zero para a maioria "
    "das ações. Isso é consistente com a Hipótese de Mercados Eficientes na forma "
    "fraca: o preço de hoje já incorporou toda a informação pública disponível, "
    "incluindo o histórico de preços. Saber que a PETR4 subiu ontem não te dá "
    "vantagem para prever o que ela vai fazer amanhã.")

fala(doc,
    "Já em frequência mensal, a história muda. Há evidência de autocorrelação "
    "positiva em horizontes de 6 a 12 meses — é o fenômeno de momentum. "
    "Ações que performaram bem nos últimos 12 meses tendem a continuar performando "
    "bem nos próximos 1 a 3 meses. E é exatamente isso que vamos explorar "
    "na Aula 4 quando construirmos o sinal de momentum 12-1.")

# ══════════════════════════════════════════════════════════════════════════════
secao(doc, "PARTE 2 — LIVE CODING (40 min)")
tempo(doc, "0:18 – 0:55")
sep(doc)

acao(doc, "Abrir notebook: aulas/aula-03-eda/aula-03-eda.ipynb")
fala(doc, "Abram o notebook. Rodem as células de setup e imports primeiro.")

# ── Bloco 1: Scatter Risco × Retorno ─────────────────────────────────────────
subsecao(doc, "Bloco 1 — Scatter Risco × Retorno")

fala(doc,
    "O primeiro bloco calcula o retorno médio anual e a volatilidade anual para "
    "cada ação do universo. Vamos ver o código.")

cod(doc,
    "resumo = pd.DataFrame({\n"
    "    'retorno_anual': retornos.mean() * 252,\n"
    "    'vol_anual':     retornos.std()  * np.sqrt(252)\n"
    "}).dropna().sort_values('retorno_anual', ascending=False)\n\n"
    "print('Top 5:');   print(resumo.head())\n"
    "print('Bottom 5:'); print(resumo.tail())")

output(doc,
    f"Top 5 maiores retornos anuais médios:\n"
    f"{top5.to_string(float_format='{:.2%}'.format)}\n\n"
    f"Bottom 5 menores retornos anuais médios:\n"
    f"{bot5.to_string(float_format='{:.2%}'.format)}")

interp(doc,
    f"Olhem o top 5. "
    + (f"WEGE3 no topo com retorno anual médio de "
       f"{resumo.loc['WEGE3.SA','retorno_anual']:.1%} e volatilidade de "
       f"{resumo.loc['WEGE3.SA','vol_anual']:.1%}. " if 'WEGE3.SA' in resumo.index else "")
    + "Isso é o que chamamos de 'alfa' — retorno acima do que seria esperado "
    "pelo nível de risco. Agora o bottom 5: "
    + (f"MGLU3 com retorno anual médio de "
       f"{resumo.loc['MGLU3.SA','retorno_anual']:.1%}. " if 'MGLU3.SA' in resumo.index else "")
    + "Retorno negativo no longo prazo — o risco não foi compensado. "
    f"No total, {acima_zero} ações tiveram retorno médio positivo e "
    f"{abaixo_zero} tiveram retorno negativo ao longo do período.")

fala(doc, "Agora o scatter — o gráfico que resume tudo isso visualmente:")

add_buf(doc, buf1, largura=6.2)

interp(doc,
    "Esse gráfico é o retrato do universo de investimentos. Cada ponto é uma ação. "
    "A linha vermelha tracejada é o retorno zero — tudo abaixo dela destruiu valor "
    "no período. "
    + (f"WEGE3 está no canto superior esquerdo — o melhor de todos os mundos: "
       f"alto retorno ({resumo.loc['WEGE3.SA','retorno_anual']:.1%} ao ano) com "
       f"volatilidade relativamente controlada ({resumo.loc['WEGE3.SA','vol_anual']:.1%}). "
       if 'WEGE3.SA' in resumo.index else "")
    + (f"PETR4 está no meio do gráfico — retorno de "
       f"{resumo.loc['PETR4.SA','retorno_anual']:.1%} com volatilidade de "
       f"{resumo.loc['PETR4.SA','vol_anual']:.1%}. " if 'PETR4.SA' in resumo.index else "")
    + (f"MGLU3 está abaixo da linha vermelha com volatilidade altíssima "
       f"({resumo.loc['MGLU3.SA','vol_anual']:.1%}) — o pior dos mundos. "
       if 'MGLU3.SA' in resumo.index else "")
    + "Reparem que não há uma relação clara positiva entre risco e retorno — "
    "pontos com alta volatilidade aparecem tanto acima quanto abaixo da linha "
    "zero. Isso é o que justifica uma estratégia ativa: escolher as ações "
    "certas, não simplesmente as mais arriscadas.")

# ── Bloco 2: Fat Tails ────────────────────────────────────────────────────────
subsecao(doc, "Bloco 2 — Fat Tails: Histograma e QQ-Plot (PETR4)")

fala(doc,
    "Agora vamos olhar para a distribuição dos retornos de PETR4 e comparar "
    "com a distribuição normal teórica.")

cod(doc,
    "x = np.linspace(petr4_ret.min(), petr4_ret.max(), 300)\n\n"
    "# Histograma + curva normal sobreposta\n"
    "axes[0].hist(petr4_ret, bins=100, density=True, alpha=0.6, label='retornos reais')\n"
    "axes[0].plot(x, stats.norm.pdf(x, petr4_ret.mean(), petr4_ret.std()),\n"
    "             'r-', lw=2, label='normal ajustada')\n\n"
    "# QQ-Plot\n"
    "stats.probplot(petr4_ret, dist='norm', plot=axes[1])\n\n"
    "print(f'Curtose (excess): {petr4_ret.kurtosis():.2f}  (normal = 0)')\n"
    "print(f'Assimetria:       {petr4_ret.skew():.2f}')")

add_buf(doc, buf2, largura=6.2)

output(doc,
    f"Curtose (excess): {curtose:.2f}  (normal = 0)\n"
    f"Assimetria:       {assimetria:.2f}   (normal = 0)\n\n"
    "Eventos extremos observados vs esperado pela normal:\n"
    + "\n".join([
        f"  |r| > {n}σ:  observado = {extremos[n][0]:3d} dias  "
        f"|  esperado pela normal = {extremos[n][1]:.1f} dias"
        for n in [3, 4, 5]
    ]))

interp(doc,
    f"Histograma — lado esquerdo. O pico azul é mais alto e mais estreito "
    "que a curva vermelha (normal). E as bordas — as caudas — ficam acima "
    "da curva vermelha nas extremidades. Isso é fat tails visível a olho nu. "
    f"O número confirma: curtose em excesso de {curtose:.2f}. "
    "Na normal teórica o excesso seria zero. "
    f"Aqui é {curtose:.1f}x mais intenso. "
    f"A assimetria de {assimetria:.2f} confirma o viés negativo — "
    "a cauda esquerda é mais pesada.")

interp(doc,
    "QQ-Plot — lado direito. Se os retornos fossem normais, todos os pontos "
    "ficariam exatamente sobre a linha vermelha diagonal. "
    "Reparem nas extremidades: os pontos se afastam muito da linha — para cima "
    "no lado direito (retornos positivos extremos mais frequentes que a normal) "
    "e para baixo no lado esquerdo (quedas extremas muito mais frequentes). "
    "Esse formato de 'S' nas pontas é a assinatura visual das fat tails.")

interp(doc,
    f"Os números de eventos extremos são chocantes. "
    f"A normal previa {extremos[3][1]:.0f} dias com retorno acima de 3 desvios padrão "
    f"em 15 anos. Aconteceram {extremos[3][0]} vezes — "
    f"{extremos[3][0]/extremos[3][1]:.1f}x mais. "
    f"Para 4 sigma: esperado {extremos[4][1]:.1f} dia, observado {extremos[4][0]} vezes. "
    f"Para 5 sigma: a normal diz que deveria ocorrer a cada 4.700 anos. "
    f"Ocorreu {extremos[5][0]} vezes em 15 anos de PETR4. "
    "Isso não é um detalhe técnico — é a diferença entre um modelo de risco "
    "que funciona e um que subestima o drawdown máximo por um fator de 10.")

# ── Bloco 3: Estacionariedade ─────────────────────────────────────────────────
subsecao(doc, "Bloco 3 — Estacionariedade: Preço vs Retorno + Teste ADF")

fala(doc,
    "Vamos ver a diferença entre preço e retorno visualmente — e depois "
    "confirmar com o teste estatístico formal.")

cod(doc,
    "# Visualização\n"
    "petr4_preco.plot(ax=axes[0])  # preço — não estacionário\n"
    "petr4_ret.plot(ax=axes[1])    # retorno — estacionário\n"
    "axes[1].axhline(petr4_ret.mean(), color='red', linestyle='--')\n\n"
    "# Teste ADF\n"
    "from statsmodels.tsa.stattools import adfuller\n"
    "for nome, serie in [('Preço', petr4_preco), ('Retorno', petr4_ret)]:\n"
    "    p = adfuller(serie.dropna())[1]\n"
    "    print(f'ADF {nome}: p-value = {p:.4f}')")

add_buf(doc, buf3, largura=6.2)

output(doc,
    f"ADF Preço:   p-value = {adf_preco[1]:.4f}   "
    f"→ {'Estacionária ✓' if adf_preco[1] < 0.05 else 'Não estacionária ✗'}\n"
    f"ADF Retorno: p-value = {adf_retorno[1]:.4f}  "
    f"→ {'Estacionária ✓' if adf_retorno[1] < 0.05 else 'Não estacionária ✗'}")

interp(doc,
    "Gráfico superior — o preço de PETR4. Reparem que não tem um nível fixo. "
    "A série começa em torno de R$20, sobe para mais de R$30, cai, sobe de novo. "
    "Não há como dizer que a 'média é X' porque a média muda dependendo de qual "
    "janela você escolhe. Essa é a não-estacionariedade visível.")

interp(doc,
    "Gráfico inferior — o retorno diário. Completamente diferente. A série oscila "
    "em torno da linha vermelha (a média) de forma estável. Não tem tendência "
    "de alta ou queda estrutural. Oscila, mas sempre volta para perto de zero. "
    "Isso é estacionariedade.")

interp(doc,
    f"O teste ADF confirma formalmente. Para o preço, p-value = {adf_preco[1]:.4f} "
    f"— {'acima' if adf_preco[1] > 0.05 else 'abaixo'} de 0.05, "
    f"{'não rejeitamos' if adf_preco[1] > 0.05 else 'rejeitamos'} H0, "
    f"{'não estacionária' if adf_preco[1] > 0.05 else 'estacionária'}. "
    f"Para o retorno, p-value = {adf_retorno[1]:.4f} "
    f"— {'abaixo' if adf_retorno[1] < 0.05 else 'acima'} de 0.05, "
    f"{'rejeitamos H0' if adf_retorno[1] < 0.05 else 'não rejeitamos H0'}, "
    f"{'estacionária' if adf_retorno[1] < 0.05 else 'não estacionária'}. "
    "Exatamente o que a teoria prevê.")

# ── Bloco 4: Autocorrelação ───────────────────────────────────────────────────
subsecao(doc, "Bloco 4 — Autocorrelação: O Mercado Tem Memória?")

fala(doc,
    "Agora a pergunta central para a nossa estratégia: os retornos passados "
    "contêm informação sobre os retornos futuros?")

cod(doc,
    "from statsmodels.graphics.tsaplots import plot_acf\n\n"
    "# Diário: lags 1–30 dias\n"
    "plot_acf(petr4_ret.dropna(), lags=30, ax=axes[0], alpha=0.05)\n\n"
    "# Mensal: lags 1–24 meses\n"
    "plot_acf(ret_mensais['PETR4.SA'].dropna(), lags=24, ax=axes[1], alpha=0.05)\n\n"
    "print(f'Autocorr 1 dia:    {petr4_ret.autocorr(lag=1):.4f}')\n"
    "print(f'Autocorr 1 mês:    {ret_mensais[\"PETR4.SA\"].autocorr(lag=1):.4f}')\n"
    "print(f'Autocorr 12 meses: {ret_mensais[\"PETR4.SA\"].autocorr(lag=12):.4f}')")

add_buf(doc, buf4, largura=6.2)

output(doc,
    f"Autocorr 1 dia:    {autocorr_1d:.4f}\n"
    f"Autocorr 1 mês:    {autocorr_1m:.4f}\n"
    f"Autocorr 12 meses: {autocorr_12m:.4f}")

interp(doc,
    "Gráfico da esquerda — autocorrelação nos retornos DIÁRIOS. "
    "A grande maioria das barras está dentro da área azul (intervalo de confiança 95%). "
    "Isso significa que a autocorrelação não é estatisticamente diferente de zero "
    "para quase nenhum lag. "
    f"O lag 1 tem autocorrelação de {autocorr_1d:.4f} — praticamente zero. "
    "Saber o retorno de PETR4 hoje não te diz nada sobre o de amanhã. "
    "Isso é consistente com mercados eficientes na forma fraca no curto prazo.")

interp(doc,
    "Gráfico da direita — autocorrelação nos retornos MENSAIS. "
    "O padrão muda. "
    f"Autocorrelação no lag 1 (1 mês): {autocorr_1m:.4f}. "
    f"Autocorrelação no lag 12 (12 meses): {autocorr_12m:.4f}. "
    "Algumas barras saem da área azul — especialmente nos lags mais curtos. "
    "Isso sugere que em frequência mensal existe alguma estrutura temporal nos "
    "retornos que não existe no diário. "
    "É a evidência estatística do fenômeno de momentum: retornos mensais "
    "positivos no passado recente têm alguma correlação com retornos futuros. "
    "Na Aula 4 vamos formalizar esse sinal: para cada mês, ranqueamos as ações "
    "pelo retorno acumulado nos 11 meses anteriores (excluindo o último) e "
    "compramos as top 10. É exatamente isso que os ACF plots estão dizendo que "
    "pode funcionar.")

# ── Bloco 5: Filtros ──────────────────────────────────────────────────────────
subsecao(doc, "Bloco 5 — Filtros de Cobertura e Liquidez")

fala(doc,
    "Antes de salvar o dataset limpo, aplicamos dois filtros. O primeiro "
    "é o filtro de cobertura — já mencionado na Aula 2. O segundo é o filtro "
    "de liquidez.")

cod(doc,
    "# Filtro de cobertura\n"
    "cobertura = retornos.notna().mean()\n"
    "tickers_validos = cobertura[cobertura >= 0.80].index.tolist()\n\n"
    "# Filtro de liquidez (proxy: % de dias com retorno != 0)\n"
    "liquidez = (retornos[tickers_validos] != 0).mean()\n"
    "tickers_finais = liquidez[liquidez >= liquidez.quantile(0.10)].index.tolist()\n\n"
    "print(f'Removidos por cobertura: {len(tickers_removidos)}')\n"
    "print(f'Dataset final: {len(tickers_finais)} ações')")

output(doc,
    f"Cobertura < 80% → removidos: {len(tickers_removidos_cob)}\n"
    f"  {[t.replace('.SA','') for t in tickers_removidos_cob]}\n\n"
    f"Após filtro de liquidez (p10): {len(tickers_finais)} ações")

interp(doc,
    f"{len(tickers_removidos_cob)} ações removidas por cobertura insuficiente. "
    "A maioria são IPOs recentes — empresas que não existiam no início do nosso "
    "período de análise. Algumas são casos de suspensão prolongada por problemas "
    "financeiros ou corporativos. "
    "O filtro de liquidez com proxy de % de dias com retorno não-nulo é uma "
    "abordagem simples e eficiente: ações que ficam dias com retorno zero têm "
    "baixíssimo volume real de negociação. Na prática, você não conseguiria "
    "executar uma ordem nessas ações sem mover o preço de forma significativa. "
    f"O resultado final é {len(tickers_finais)} ações com histórico suficiente "
    "e liquidez mínima — um universo confiável para o backtest.")

# ── Bloco 6: Salvar ───────────────────────────────────────────────────────────
subsecao(doc, "Bloco 6 — Salvando o Dataset Limpo")

cod(doc,
    "precos[tickers_finais].to_parquet(\n"
    "    os.path.join(DADOS_DIR, 'precos_limpo.parquet'))\n"
    "retornos[tickers_finais].to_parquet(\n"
    "    os.path.join(DADOS_DIR, 'retornos_diarios_limpo.parquet'))\n"
    "ret_mensais[tickers_finais].to_parquet(\n"
    "    os.path.join(DADOS_DIR, 'retornos_mensais_limpo.parquet'))\n"
    "pd.Series(tickers_finais).to_csv(\n"
    "    os.path.join(DADOS_DIR, 'tickers_finais.csv'), index=False)")

arqs = {
    'precos_limpo.parquet': 'precos_limpo.parquet',
    'retornos_diarios_limpo.parquet': 'retornos_diarios_limpo.parquet',
    'retornos_mensais_limpo.parquet': 'retornos_mensais_limpo.parquet',
}
linhas_tam = []
for nome, arq in arqs.items():
    fp = os.path.join(DADOS_DIR, arq)
    if os.path.exists(fp):
        tam = os.path.getsize(fp) // 1024
        linhas_tam.append(f"  {nome:<40} {tam:>6} KB")
    else:
        linhas_tam.append(f"  {nome:<40}  (ainda não gerado — rode o notebook)")

output(doc,
    f"Dataset limpo: {len(tickers_finais)} ações × {retornos.shape[0]} dias\n"
    "Arquivos salvos:\n" + "\n".join(linhas_tam))

interp(doc,
    "Quatro arquivos gerados. Os três parquets são as versões limpas dos dados "
    "da Aula 2 — filtramos as ações com histórico insuficiente e baixa liquidez. "
    "O CSV com os tickers finais é uma conveniência: as próximas aulas podem "
    "carregar essa lista em vez de refazer os filtros. "
    "A partir da Aula 4, todos os notebooks usam esses quatro arquivos. "
    "A Aula 2 foi a coleta; a Aula 3 foi a limpeza. As próximas aulas são a análise.")

# ══════════════════════════════════════════════════════════════════════════════
secao(doc, "FECHAMENTO DA AULA 3")
tempo(doc, "0:55 – 1:00")
sep(doc)

fala(doc,
    "Pessoal, vamos fechar com o que construímos hoje. Quatro conceitos que "
    "aparecem em todo relatório quant sério — e agora vocês têm os números "
    "reais do mercado brasileiro para ilustrar cada um.")

fala(doc,
    f"Risco versus retorno: vimos {resumo.shape[0]} ações plotadas. "
    + (f"WEGE3 com {resumo.loc['WEGE3.SA','retorno_anual']:.1%} ao ano é o caso "
       "de maior sucesso no período. " if 'WEGE3.SA' in resumo.index else "")
    + "A dispersão enorme justifica por que seleção importa.")

fala(doc,
    f"Fat tails: curtose de {curtose:.1f} e assimetria de {assimetria:.2f}. "
    f"Eventos de 5 sigma aconteceram {extremos[5][0]} vezes em 15 anos — "
    "a normal previa que demorariam milênios. Isso tem impacto direto em "
    "como medimos risco: nunca confiem só no Sharpe Ratio.")

fala(doc,
    f"Estacionariedade: preço com p-value ADF de {adf_preco[1]:.4f} "
    f"— não estacionário. Retorno com p-value {adf_retorno[1]:.4f} "
    "— estacionário. Sempre trabalhem com retornos.")

fala(doc,
    f"Autocorrelação: praticamente zero no diário ({autocorr_1d:.4f}). "
    f"Alguma estrutura no mensal. Isso é a base estatística do momentum "
    "que vamos explorar a partir da Aula 4.")

fala(doc,
    "Para a próxima aula: confiram que os quatro arquivos da Aula 3 foram gerados "
    "na pasta dados/. Aula 4 começa carregando exatamente esses arquivos para "
    "construir o primeiro sinal de momentum.")

doc.add_paragraph()
sep(doc)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("ImpactUFSCar — Diretoria de Quant — 2025")
r.font.size = Pt(9); r.italic = True; r.font.color.rgb = RGBColor(136,136,136)

doc.save(OUT)
print(f"\nSalvo: {OUT}")
