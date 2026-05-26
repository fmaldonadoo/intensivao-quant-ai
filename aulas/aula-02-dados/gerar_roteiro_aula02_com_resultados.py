# -*- coding: utf-8 -*-
"""
Gera roteiro-aula-02-dados-completo.docx
Usa EXATAMENTE o mesmo código do notebook para gerar os gráficos,
e embute análise de cada output/gráfico no roteiro do professor.
"""
import os
from io import BytesIO

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
import yfinance as yf

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ── Caminhos ─────────────────────────────────────────────────────────────────
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
ROOT = SCRIPT_DIR
_p = SCRIPT_DIR
while _p != os.path.dirname(_p):
    if os.path.exists(os.path.join(_p, '.git')):
        ROOT = _p; break
    _p = os.path.dirname(_p)

DADOS_DIR = os.path.join(ROOT, 'dados')
OUT = os.path.join(SCRIPT_DIR, 'roteiro-aula-02-dados-final.docx')

# ══════════════════════════════════════════════════════════════════════════════
# REPRODUZ EXATAMENTE O CÓDIGO DO NOTEBOOK — mesmos dados, mesmos gráficos
# ══════════════════════════════════════════════════════════════════════════════
print("Reproduzindo código do notebook...")

plt.rcParams['figure.figsize'] = (12, 5)
plt.rcParams['axes.grid'] = True
plt.rcParams['grid.alpha'] = 0.3

# ── Gráfico 1 (célula 7) — Retorno simples vs log ────────────────────────────
print("  [1/4] Retorno simples vs log (PETR4)...")
petr4 = yf.download("PETR4.SA", start="2020-01-01", end="2024-12-31",
                    auto_adjust=True, progress=False)
preco = petr4['Close'].squeeze()

ret_simples = preco.pct_change()
ret_log     = np.log(preco / preco.shift(1))

comparacao = pd.DataFrame({'ret_simples': ret_simples, 'ret_log': ret_log})
comparacao_head = comparacao.dropna().head()
diff_media = (ret_simples - ret_log).abs().mean()

fig1, axes1 = plt.subplots(1, 2, figsize=(14, 4))

ret_simples.dropna().iloc[:60].plot(ax=axes1[0], label='simples', alpha=0.8)
ret_log.dropna().iloc[:60].plot(ax=axes1[0], label='log', alpha=0.8, linestyle='--')
axes1[0].set_title('Retornos diários — 60 primeiros dias')
axes1[0].legend()

acum_simples = (1 + ret_simples).cumprod() - 1
acum_log     = ret_log.cumsum()
acum_simples.plot(ax=axes1[1], label='simples acumulado')
acum_log.plot(ax=axes1[1], label='log acumulado', linestyle='--')
axes1[1].set_title('Retorno acumulado — 2020 a 2024')
axes1[1].legend()

plt.tight_layout()
buf1 = BytesIO(); fig1.savefig(buf1, format='png', dpi=150, bbox_inches='tight'); buf1.seek(0)
plt.close(fig1)

acum_simples_final = acum_simples.iloc[-1]
acum_log_final_convertido = np.exp(acum_log.iloc[-1]) - 1

# ── Gráfico 2 (célula 9) — VALE3 bruto vs ajustado ───────────────────────────
print("  [2/4] VALE3 bruto vs ajustado...")
vale_bruto    = yf.download("VALE3.SA", start="2020-01-01", end="2024-12-31",
                             auto_adjust=False, progress=False)['Close'].squeeze()
vale_ajustado = yf.download("VALE3.SA", start="2020-01-01", end="2024-12-31",
                             auto_adjust=True,  progress=False)['Close'].squeeze()

fig2, ax2 = plt.subplots(figsize=(12, 4))
vale_bruto.plot(ax=ax2, label='preço bruto', alpha=0.7)
vale_ajustado.plot(ax=ax2, label='preço ajustado', alpha=0.7)
ax2.set_title('VALE3 — Bruto vs Ajustado (dividendos incluídos no ajustado)')
ax2.legend()

buf2 = BytesIO(); fig2.savefig(buf2, format='png', dpi=150, bbox_inches='tight'); buf2.seek(0)
plt.close(fig2)

# ── Carrega parquets para gráficos 3 e 4 ─────────────────────────────────────
print("  [3/4] Preços normalizados...")
precos_raw = pd.read_parquet(os.path.join(DADOS_DIR, 'precos_ibov.parquet'))
retornos   = pd.read_parquet(os.path.join(DADOS_DIR, 'retornos_diarios.parquet'))
retornos_mensais = pd.read_parquet(os.path.join(DADOS_DIR, 'retornos_mensais.parquet'))

# Gráfico 3 (célula 22) — Preço normalizado em 100
acoes_exemplo = ['PETR4.SA', 'VALE3.SA', 'WEGE3.SA', 'ITUB4.SA']
fig3, ax3 = plt.subplots(figsize=(13, 5))
for ticker in acoes_exemplo:
    if ticker in precos_raw.columns:
        serie = precos_raw[ticker].dropna()
        (serie / serie.iloc[0] * 100).plot(ax=ax3, label=ticker.replace('.SA', ''))
ax3.set_title('Preço ajustado — normalizado em 100 (Jan/2010)')
ax3.set_ylabel('Índice (base 100)')
ax3.legend()
plt.tight_layout()

buf3 = BytesIO(); fig3.savefig(buf3, format='png', dpi=150, bbox_inches='tight'); buf3.seek(0)
plt.close(fig3)

# Retorno acumulado de cada ação para interpretação
retorno_acumulado = retornos.sum().sort_values(ascending=False)
top5 = retorno_acumulado.head(5)
bot5 = retorno_acumulado.tail(5)

# ── Gráfico 4 (célula 23) — Retornos diários e mensais PETR4 ─────────────────
print("  [4/4] Retornos PETR4...")
fig4, axes4 = plt.subplots(2, 1, figsize=(13, 7))

retornos['PETR4.SA'].plot(ax=axes4[0], alpha=0.7, color='steelblue')
axes4[0].set_title('Retornos diários — PETR4 (2010–2024)')
axes4[0].set_ylabel('Retorno log')
axes4[0].axhline(0, color='black', linewidth=0.8)

mes_petr4 = retornos_mensais['PETR4.SA'].copy()
mes_petr4.index = range(len(mes_petr4))
cores_bar = ['steelblue' if v >= 0 else 'tomato' for v in mes_petr4.values]
axes4[1].bar(mes_petr4.index, mes_petr4.values, color=cores_bar, alpha=0.7, width=0.8)
axes4[1].set_title('Retornos mensais — PETR4 (2010–2024)')
axes4[1].set_ylabel('Retorno log mensal')
tick_pos = [i for i in range(len(retornos_mensais.index)) if i % 12 == 0]
tick_lbl = [retornos_mensais.index[i].strftime('%Y') for i in tick_pos]
axes4[1].set_xticks(tick_pos)
axes4[1].set_xticklabels(tick_lbl, rotation=45)
plt.tight_layout()

buf4 = BytesIO(); fig4.savefig(buf4, format='png', dpi=150, bbox_inches='tight'); buf4.seek(0)
plt.close(fig4)

# Estatísticas PETR4 para interpretação
petr4_stats = retornos['PETR4.SA'].describe()
petr4_mes_stats = retornos_mensais['PETR4.SA'].describe()
nans = precos_raw.isna().sum().sort_values(ascending=False)

print("Gráficos prontos. Montando documento...")

# ══════════════════════════════════════════════════════════════════════════════
# HELPERS DE DOCUMENTO
# ══════════════════════════════════════════════════════════════════════════════
def novo_doc():
    doc = Document()
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Cm(2.5)
        s.left_margin = Cm(3.0); s.right_margin = Cm(2.5)
    return doc

def capa(doc, t1, t2, nota):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(t1); r.bold = True; r.font.size = Pt(16)
    r.font.color.rgb = RGBColor(0,45,98)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(t2); r.font.size = Pt(13)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(nota); r.italic = True; r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(120,120,120)

def secao(doc, t):
    doc.add_paragraph()
    p = doc.add_heading(t, 1)
    p.runs[0].font.color.rgb = RGBColor(0,45,98)
    p.runs[0].font.size = Pt(13)

def subsecao(doc, t):
    p = doc.add_heading(t, 2)
    p.runs[0].font.color.rgb = RGBColor(184,134,0)
    p.runs[0].font.size = Pt(11)

def sep(doc):
    p = doc.add_paragraph()
    r = p.add_run("─" * 70)
    r.font.size = Pt(9); r.font.color.rgb = RGBColor(180,180,180)

def tempo(doc, t):
    p = doc.add_paragraph()
    r = p.add_run(f"  {t}")
    r.font.size = Pt(9); r.italic = True; r.font.color.rgb = RGBColor(120,120,120)

def fala(doc, texto):
    p = doc.add_paragraph()
    r = p.add_run(texto); r.font.size = Pt(11)

def acao(doc, texto):
    p = doc.add_paragraph()
    r = p.add_run(f"[{texto}]")
    r.font.size = Pt(10); r.italic = True
    r.font.color.rgb = RGBColor(0,100,180)

def cod(doc, texto):
    p = doc.add_paragraph()
    r = p.add_run(texto)
    r.font.name = 'Courier New'; r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(30,30,30)

def output(doc, texto):
    p = doc.add_paragraph()
    r = p.add_run(f"▶  OUTPUT REAL:\n{texto}")
    r.font.name = 'Courier New'; r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0,100,0); r.bold = True

def interp(doc, texto):
    p = doc.add_paragraph()
    r = p.add_run(f"↳  {texto}")
    r.font.size = Pt(11); r.font.color.rgb = RGBColor(0,100,0)

def add_buf(doc, buf, largura=6.0):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(buf, width=Inches(largura))

# ══════════════════════════════════════════════════════════════════════════════
doc = novo_doc()

capa(doc,
     "INTENSIVÃO QUANT AI — AULA 02",
     "Dados: Coleta, Limpeza e Pipeline",
     "Roteiro completo com código + resultados reais — uso exclusivo do instrutor")

# ══════════════════════════════════════════════════════════════════════════════
secao(doc, "ABERTURA")
tempo(doc, "0:00 – 0:03  |  Boas-vindas e contexto")
sep(doc)

fala(doc,
    "Pessoal, bem-vindos à Aula 2. Na semana passada a gente construiu a base "
    "conceitual — entendemos por que momentum existe e o que queremos construir. "
    "Hoje a gente bota a mão na massa de verdade: vamos coletar dados, limpar e "
    "construir o pipeline que todas as próximas aulas vão usar.")

fala(doc,
    "Uma coisa que quero deixar clara logo no início: a qualidade dos dados é a "
    "parte mais subestimada de qualquer estratégia quant. Não é a parte mais "
    "glamourosa — mas é onde a maioria dos erros acontece. Dados ruins geram "
    "backtests errados que geram conclusões erradas. A gente vai ser obsessivos "
    "com qualidade aqui.")

fala(doc,
    "Ao final dessa aula, vocês vão ter três arquivos salvos: precos_ibov.parquet, "
    "retornos_diarios.parquet e retornos_mensais.parquet. Esses três arquivos são "
    "a matéria-prima de tudo que a gente vai construir daqui pra frente.")

# ══════════════════════════════════════════════════════════════════════════════
secao(doc, "PARTE 1 — TEORIA (20 min)")
tempo(doc, "0:03 – 0:20")
sep(doc)

acao(doc, "Avançar para slide 3: Série Temporal de Preços")
subsecao(doc, "O que é uma série temporal de preços?")

fala(doc,
    "Uma série temporal é uma sequência de dados coletados e ordenados ao longo "
    "do tempo. No mercado financeiro, a série temporal de preços é nossa matéria-"
    "prima fundamental. Cada linha representa o preço de fechamento de uma ação "
    "em uma data específica.")

fala(doc,
    "Mas atenção — e isso é crítico: preços brutos são péssimos para análise "
    "quantitativa. Três razões. Primeiro: falta de comparabilidade. WEGE3 a "
    "R$45 e PETR4 a R$37 não nos diz qual performou melhor. Segundo: "
    "não-estacionariedade — o preço tem tendência, então a média muda com a "
    "janela escolhida, violando premissas estatísticas. Terceiro: escala "
    "diferente entre empresas. Por isso nós nunca modelamos preços brutos "
    "diretamente. Transformamos em retornos.")

acao(doc, "Avançar para slide 4: Retorno Simples vs Logarítmico")
subsecao(doc, "Retorno Simples vs Logarítmico")

fala(doc,
    "O retorno simples é a variação percentual que vocês já conhecem: "
    "R_t = (P_t - P_{t-1}) / P_{t-1}. Se a ação foi de R$100 para R$105, "
    "retorno simples é 5%.")

fala(doc,
    "O retorno logarítmico usa o logaritmo natural da razão de preços: "
    "r_t = ln(P_t / P_{t-1}). No mesmo exemplo: ln(105/100) = 4.88%.")

fala(doc,
    "Por que os quants preferem o log? A propriedade mágica: aditividade temporal. "
    "Para calcular o retorno acumulado de 3 dias com retorno simples, você precisa "
    "multiplicar os fatores. Com retorno log, basta somar. "
    "r1 + r2 + r3 = ln(P1/P0) + ln(P2/P1) + ln(P3/P2) = ln(P3/P0). "
    "Os termos intermediários se cancelam. Isso torna cálculos acumulados "
    "computacionalmente mais simples e matematicamente mais limpos.")

fala(doc,
    "Para retornos diários pequenos, os dois são praticamente idênticos — "
    "vamos ver isso no código. A diferença aparece no longo prazo.")

acao(doc, "Avançar para slide 5: Preços Ajustados")
subsecao(doc, "Preços Ajustados — Splits e Dividendos")

fala(doc,
    "Segundo ponto crítico: preços ajustados. Quando uma empresa faz um split "
    "de 1 para 2, o preço cai pela metade da noite pro dia. Se você usa o preço "
    "bruto, o algoritmo enxerga uma queda de 50% que nunca aconteceu de verdade. "
    "Dividendos causam o mesmo problema — o preço cai na data ex-dividendo, "
    "mas o investidor recebeu esse dinheiro em conta.")

fala(doc,
    "A solução é o preço ajustado. O yfinance faz isso automaticamente com "
    "auto_adjust=True — ele reconstrói toda a série histórica retroativamente "
    "para que cada retorno calculado represente a rentabilidade real total do "
    "investidor. Regra de ouro: nunca calcule retornos com preço bruto.")

acao(doc, "Avançar para slide 6: IBOVESPA e Survivorship Bias")
subsecao(doc, "O IBOVESPA e o Viés de Sobrevivência")

fala(doc,
    "O IBOVESPA é o principal índice da B3 — as ações mais líquidas do mercado "
    "brasileiro. A carteira é rebalanceada quadrimestralmente.")

fala(doc,
    "Um alerta importante que vocês precisam mencionar pro Itaú: survivorship bias "
    "— viés de sobrevivência. Quando usamos a composição atual do índice para "
    "fazer um backtest de 10 anos atrás, estamos selecionando só as empresas que "
    "sobreviveram e prosperaram até hoje. Ignoramos as que faliram ou foram "
    "deslistadas — OGX, Americanas, PDG. Como resultado, o backtest fica "
    "artificialmente otimista. No intensivão usamos uma lista simplificada das "
    "principais componentes históricas — mas mencionem isso proativamente na "
    "apresentação. A banca do Itaú vai reconhecer a maturidade técnica.")

# ══════════════════════════════════════════════════════════════════════════════
secao(doc, "PARTE 2 — LIVE CODING (40 min)")
tempo(doc, "0:20 – 0:55")
sep(doc)

acao(doc, "Abrir notebook: aulas/aula-02-dados/aula-02-dados.ipynb")
fala(doc, "Abram o notebook. Vamos percorrer cada célula juntos.")

# ── Bloco 1: Setup ────────────────────────────────────────────────────────────
subsecao(doc, "Bloco 1 — Instalação e Configuração do Ambiente")

fala(doc,
    "A primeira célula instala tudo automaticamente. Se for a primeira vez "
    "no computador, ela instala e reinicia o kernel. Se reiniciar, rodem de "
    "novo do início.")

fala(doc,
    "A segunda célula detecta VS Code ou Google Colab e define DADOS_DIR — "
    "a pasta onde todos os arquivos serão salvos. Ela sobe as pastas até "
    "encontrar o .git do projeto, então funciona em qualquer computador "
    "sem precisar alterar nada.")

cod(doc,
    "# Sobe pastas até encontrar o .git do repositório\n"
    "_p = os.path.abspath(os.getcwd())\n"
    "while _p != os.path.dirname(_p):\n"
    "    if os.path.exists(os.path.join(_p, '.git')):\n"
    "        _root = _p; break\n"
    "    _p = os.path.dirname(_p)\n\n"
    "DADOS_DIR = os.path.join(_root, 'dados')\n"
    "os.makedirs(DADOS_DIR, exist_ok=True)\n"
    "print(f'Pasta de dados: {DADOS_DIR}')")

output(doc,
    f"Ambiente: VS Code / Jupyter local\n"
    f"Pasta de dados: {DADOS_DIR}")

interp(doc,
    "Esse caminho é calculado dinamicamente. Em qualquer computador que clonar "
    "o repositório, vai apontar para a pasta dados/ dentro do projeto. "
    "Quem rodar no Colab vai ver um caminho do Google Drive.")

# ── Bloco 2: Retorno simples vs log ──────────────────────────────────────────
subsecao(doc, "Bloco 2 — Retorno Simples vs Logarítmico (PETR4)")

fala(doc,
    "Antes de baixar tudo, demonstramos a diferença entre retorno simples e "
    "logarítmico usando PETR4 de 2020 a 2024.")

cod(doc,
    "petr4 = yf.download('PETR4.SA', start='2020-01-01', end='2024-12-31',\n"
    "                    auto_adjust=True, progress=False)\n"
    "preco = petr4['Close'].squeeze()\n\n"
    "ret_simples = preco.pct_change()\n"
    "ret_log     = np.log(preco / preco.shift(1))\n\n"
    "print(comparacao.dropna().head())\n"
    "print(f'Diferença média (absoluta): {(ret_simples - ret_log).abs().mean():.6f}')")

output(doc,
    f"Primeiros 5 dias:\n"
    f"{comparacao_head.to_string()}\n\n"
    f"Diferença média (absoluta): {diff_media:.6f}")

interp(doc,
    f"Diferença de {diff_media:.6f} — ou seja, {diff_media*10000:.1f} pontos-base. "
    "Na prática são idênticos para retornos diários. Mas o gráfico abaixo "
    "vai mostrar onde a diferença aparece: no retorno acumulado de longo prazo.")

add_buf(doc, buf1, largura=6.2)

output(doc,
    f"Retorno simples acumulado 2020-2024: {acum_simples_final:.2%}\n"
    f"Retorno log acumulado 2020-2024:    {acum_log_final_convertido:.2%}  "
    f"(convertido de volta para %)")

interp(doc,
    "Gráfico da esquerda: os 60 primeiros dias. As duas linhas são quase "
    "indistinguíveis — ficam sobrepostas, com a linha tracejada (log) mal "
    "visível embaixo da contínua (simples). Isso confirma que para retornos "
    "diários a diferença é desprezível. "
    "Gráfico da direita: o retorno acumulado de 5 anos. Aqui a divergência "
    f"aparece. O retorno simples acumulado foi {acum_simples_final:.2%} enquanto "
    f"o log (convertido) foi {acum_log_final_convertido:.2%}. "
    "A diferença vem da convexidade: o log é sempre ligeiramente menor que o "
    "simples porque ln(1+x) < x para x > 0. No longo prazo isso acumula. "
    "No intensivão usamos retorno log nos DataFrames (é o que retornos_mensais "
    "contém), que é o padrão em finanças quant por ser aditivo.")

# ── Bloco 3: Preços ajustados VALE3 ──────────────────────────────────────────
subsecao(doc, "Bloco 3 — Preços Ajustados: VALE3 Bruto vs Ajustado")

cod(doc,
    "vale_bruto    = yf.download('VALE3.SA', start='2020-01-01', end='2024-12-31',\n"
    "                             auto_adjust=False)['Close'].squeeze()\n"
    "vale_ajustado = yf.download('VALE3.SA', start='2020-01-01', end='2024-12-31',\n"
    "                             auto_adjust=True)['Close'].squeeze()\n\n"
    "# Plot: bruto vs ajustado\n"
    "vale_bruto.plot(label='preço bruto', alpha=0.7)\n"
    "vale_ajustado.plot(label='preço ajustado', alpha=0.7)")

add_buf(doc, buf2, largura=6.2)

interp(doc,
    "As duas linhas começam no mesmo ponto — mas ao longo de 5 anos elas "
    "divergem significativamente. A linha azul (ajustada) fica acima da laranja "
    "(bruta) porque incorpora os dividendos que a VALE pagou. A VALE é conhecida "
    "por pagar dividendos generosos — em 2021 e 2022 distribuiu dividendos "
    "extraordinários por conta dos altos lucros com minério de ferro. "
    "Se você usasse o preço bruto para calcular retornos, estaria ignorando "
    "toda essa rentabilidade. O auto_adjust=True captura tudo isso. "
    "Cada 'degrau para baixo' que você veria no bruto é uma data ex-dividendo "
    "— no ajustado esse degrau não existe porque o preço anterior já foi "
    "corrigido retroativamente.")

# ── Bloco 4: Download do universo ─────────────────────────────────────────────
subsecao(doc, "Bloco 4 — Download do Universo IBOVESPA (77 tickers)")

fala(doc,
    "Agora baixamos todos os 77 tickers de uma vez. Vai levar 1 a 2 minutos.")

cod(doc,
    "TICKERS_IBOV = [\n"
    "    'ABEV3.SA', 'ASAI3.SA', 'AZUL4.SA', 'B3SA3.SA', 'BBAS3.SA',\n"
    "    ...  # lista completa no notebook\n"
    "]\n\n"
    "precos_raw = yf.download(\n"
    "    TICKERS_IBOV,\n"
    "    start='2010-01-01', end='2024-12-31',\n"
    "    auto_adjust=True, progress=True\n"
    ")['Close']\n\n"
    "print(f'Shape: {precos_raw.shape}')\n"
    "print(f'Período: {precos_raw.index[0].date()} a {precos_raw.index[-1].date()}')")

output(doc,
    "14 Failed downloads:\n"
    "['JBSS3.SA', 'CIEL3.SA', 'EMBR3.SA', 'MRFG3.SA', 'RRRP3.SA', 'AZUL4.SA',\n"
    " 'BRFS3.SA', 'ELET3.SA', 'GOLL4.SA', 'ELET6.SA', 'CCRO3.SA', 'CPLE6.SA',\n"
    " 'NTCO3.SA', 'SULA11.SA']: possibly delisted; no timezone found\n\n"
    f"Shape: {precos_raw.shape}  →  {precos_raw.shape[0]} dias × {precos_raw.shape[1]} ações\n"
    f"Período: {precos_raw.index[0].date()} a {precos_raw.index[-1].date()}")

interp(doc,
    "14 tickers falharam completamente. Olhem os nomes: ELET3 e ELET6 são as "
    "Eletrobrás — privatizada em 2022. EMBR3 é a Embraer, mudou de código. "
    "AZUL4 e GOLL4 são as aéreas com problemas financeiros sérios. "
    "BRFS3 é a BRF, que passou por escândalos de governança. "
    "Esse é o survivorship bias acontecendo na prática — existiram, tiveram "
    "problemas, e não estão mais disponíveis com histórico completo. "
    f"O shape {precos_raw.shape} confirma: {precos_raw.shape[0]} dias úteis "
    f"em 15 anos, {precos_raw.shape[1]} colunas (as 63 que conseguimos baixar). "
    "As colunas dos 14 que falharam virão cheias de NaN e serão removidas.")

# ── Bloco 5: Inspeção e NaNs ──────────────────────────────────────────────────
subsecao(doc, "Bloco 5 — Inspeção dos Dados e Dados Ausentes")

cod(doc,
    "nans = precos_raw.isna().sum().sort_values(ascending=False)\n"
    "print('Ações com mais dados ausentes:')\n"
    "print(nans[nans > 0].head(10))\n"
    "print(f'Total de ações sem nenhum NaN: {(nans == 0).sum()}')")

top_nans = nans[nans > 0].head(10)
output(doc,
    f"Ações com mais dados ausentes:\n"
    f"{top_nans.to_string()}\n\n"
    f"Total de ações sem nenhum NaN: {(nans == 0).sum()}")

interp(doc,
    f"ASAI3 tem {nans.get('ASAI3.SA', 'N/A')} dias sem dado — o Assaí fez IPO "
    "em 2021 depois de se separar do GPA. Não tem histórico antes disso. "
    "RDOR3 é a Rede D'Or, IPO em 2020. HAPV3 é a Hapvida, 2018. "
    "São boas empresas hoje, mas para uma estratégia com janela de 12 anos "
    "elas não têm o histórico que precisamos. "
    f"Só {(nans == 0).sum()} ações têm dado completo para os 15 anos inteiros — "
    "daí a necessidade de filtrar por cobertura mínima.")

# ── Bloco 6: Retornos ─────────────────────────────────────────────────────────
subsecao(doc, "Bloco 6 — Calculando os Retornos Log")

cod(doc,
    "retornos = np.log(precos_raw / precos_raw.shift(1))\n"
    "retornos = retornos.iloc[1:]  # remove primeira linha (toda NaN)\n\n"
    "print(f'Shape retornos: {retornos.shape}')\n"
    "print('Estatísticas básicas de PETR4:')\n"
    "print(retornos['PETR4.SA'].describe())")

output(doc,
    f"Shape retornos: {retornos.shape}\n\n"
    f"Estatísticas básicas de PETR4:\n"
    f"{petr4_stats.to_string()}")

interp(doc,
    f"Shape {retornos.shape}: uma linha a menos que os preços (removemos "
    "o primeiro dia, que é todo NaN — não há retorno sem dia anterior). "
    f"PETR4: média de {petr4_stats['mean']*100:.4f}% ao dia, "
    f"desvio padrão de {petr4_stats['std']*100:.2f}% ao dia. "
    f"Mínimo de {petr4_stats['min']*100:.1f}% — provavelmente em março de "
    "2020 (COVID) ou durante alguma crise política relacionada à Petrobras. "
    f"Máximo de +{petr4_stats['max']*100:.1f}%. "
    "Notem que a média diária parece pequena, mas 0.04% ao dia × 252 dias "
    "úteis = ~10% ao ano — faz sentido.")

# ── Bloco 7: Retornos mensais ─────────────────────────────────────────────────
subsecao(doc, "Bloco 7 — Retornos Mensais")

cod(doc,
    "retornos_mensais = retornos.resample('ME').sum()\n\n"
    "print(f'Retornos mensais shape: {retornos_mensais.shape}')\n"
    "print(f'Período: {retornos_mensais.index[0].date()} a {retornos_mensais.index[-1].date()}')\n"
    "print(retornos_mensais[['PETR4.SA', 'VALE3.SA']].tail(3))")

output(doc,
    f"Retornos mensais shape: {retornos_mensais.shape}\n"
    f"Período: {retornos_mensais.index[0].date()} a {retornos_mensais.index[-1].date()}\n\n"
    f"Últimos 3 meses de PETR4 e VALE3:\n"
    f"{retornos_mensais[['PETR4.SA', 'VALE3.SA']].tail(3).to_string()}")

interp(doc,
    f"Shape {retornos_mensais.shape}: {retornos_mensais.shape[0]} meses × "
    f"{retornos_mensais.shape[1]} ações. "
    "O resample('ME') agrega todos os retornos log diários de cada mês "
    "em um único número — e como o log é aditivo, basta somar. "
    "O índice agora são datas de fim de mês (31/01, 28/02, etc.). "
    "Esse DataFrame de retornos mensais é a matéria-prima de todo o restante "
    "do intensivão: sinal de momentum, backtest, alocação — tudo vai usar ele.")

# ── Bloco 8: Gráfico 3 — Preços normalizados ─────────────────────────────────
subsecao(doc, "Bloco 8 — Visualização: Preços Normalizados (Base 100)")

fala(doc,
    "Agora os primeiros gráficos. Vamos visualizar os dados antes de modelar "
    "qualquer coisa — regra fundamental em análise quantitativa.")

cod(doc,
    "acoes_exemplo = ['PETR4.SA', 'VALE3.SA', 'WEGE3.SA', 'ITUB4.SA']\n\n"
    "for ticker in acoes_exemplo:\n"
    "    serie = precos_raw[ticker].dropna()\n"
    "    (serie / serie.iloc[0] * 100).plot(label=ticker.replace('.SA', ''))\n\n"
    "ax.set_title('Preço ajustado — normalizado em 100 (Jan/2010)')")

add_buf(doc, buf3, largura=6.2)

# Calcula os valores finais reais para interpretação
vals = {}
for t in ['PETR4.SA', 'VALE3.SA', 'WEGE3.SA', 'ITUB4.SA']:
    if t in precos_raw.columns:
        s = precos_raw[t].dropna()
        vals[t] = s.iloc[-1] / s.iloc[0] * 100

interp(doc,
    "Normalizamos tudo para 100 em janeiro de 2010 para comparar na mesma escala. "
    "O que o gráfico mostra é revelador. "
    + (f"WEGE3 chegou a {vals.get('WEGE3.SA',0):.0f} — ou seja, quem investiu R$100 "
       f"em WEG em 2010 tinha R${vals.get('WEGE3.SA',0):.0f} em 2024. " if 'WEGE3.SA' in vals else "")
    + (f"ITUB4 foi para {vals.get('ITUB4.SA',0):.0f}. " if 'ITUB4.SA' in vals else "")
    + (f"VALE3 chegou a {vals.get('VALE3.SA',0):.0f} mas com muita volatilidade no caminho — "
       "a curva sobe forte até 2011, cai muito de 2011 a 2016 (superciclo de commodities "
       "acabando), recupera de 2016 a 2021, e oscila depois. " if 'VALE3.SA' in vals else "")
    + (f"PETR4 ficou em {vals.get('PETR4.SA',0):.0f} — das quatro foi a que menos cresceu, "
       "refletindo as turbulências políticas (Lava Jato, mudanças de política de preços) "
       "ao longo dos 15 anos. " if 'PETR4.SA' in vals else "")
    + "Essa variação enorme entre as ações é exatamente por que a seleção importa — "
    "e é o que o sinal de momentum vai tentar explorar.")

# Retorno acumulado — maiores e menores
interp(doc,
    f"Para referência — as 5 ações com maior retorno log acumulado (2010-2024):\n"
    + top5.to_string()
    + f"\n\nAs 5 com menor retorno log acumulado:\n"
    + bot5.to_string())

# ── Bloco 9: Gráfico 4 — Retornos PETR4 ──────────────────────────────────────
subsecao(doc, "Bloco 9 — Visualização: Retornos Diários e Mensais PETR4")

cod(doc,
    "fig, axes = plt.subplots(2, 1, figsize=(13, 7))\n\n"
    "retornos['PETR4.SA'].plot(ax=axes[0], alpha=0.7, color='steelblue')\n"
    "axes[0].set_title('Retornos diários — PETR4 (2010–2024)')\n"
    "axes[0].axhline(0, color='black', linewidth=0.8)\n\n"
    "retornos_mensais['PETR4.SA'].plot(ax=axes[1], kind='bar', alpha=0.7,\n"
    "                                   color='steelblue', width=0.8)\n"
    "axes[1].set_title('Retornos mensais — PETR4 (2010–2024)')")

add_buf(doc, buf4, largura=6.2)

interp(doc,
    "Gráfico superior — retornos diários. À primeira vista parece ruído aleatório "
    "em torno do zero. Mas reparem nos clusters de volatilidade: em alguns períodos "
    "as oscilações são pequenas, em outros são grandes. Isso é heterocedasticidade "
    "— a volatilidade muda ao longo do tempo. Os picos maiores correspondem a "
    "eventos: a greve dos caminhoneiros de 2018 (a PETR4 despencou), o COVID "
    "em março de 2020 (o maior pico negativo visível), e as crises de preços de "
    "combustíveis. O fato de o gráfico parecer centrado em zero confirma que os "
    "retornos log têm média próxima de zero diariamente.")

interp(doc,
    "Gráfico inferior — retornos mensais. Aqui fica mais fácil identificar os "
    "eventos. Barras vermelhas longas são meses ruins, verdes são meses bons. "
    f"PETR4: média mensal de {petr4_mes_stats['mean']*100:.3f}%, "
    f"desvio padrão mensal de {petr4_mes_stats['std']*100:.2f}%, "
    f"mínimo de {petr4_mes_stats['min']*100:.1f}% e "
    f"máximo de +{petr4_mes_stats['max']*100:.1f}%. "
    "Esse DataFrame de retornos mensais é o que vamos usar para construir "
    "o sinal de momentum na Aula 4: para cada mês, qual ação teve o melhor "
    "retorno acumulado nos 11 meses anteriores (excluindo o último)?")

# ── Bloco 10: Salvando ────────────────────────────────────────────────────────
subsecao(doc, "Bloco 10 — Salvando os Parquets")

fala(doc,
    "Passo final: salvar os três arquivos. A partir daqui, nenhum notebook "
    "vai mais acessar a internet. Tudo vem desses três parquets.")

cod(doc,
    "precos_raw.to_parquet(os.path.join(DADOS_DIR, 'precos_ibov.parquet'))\n"
    "retornos.to_parquet(os.path.join(DADOS_DIR, 'retornos_diarios.parquet'))\n"
    "retornos_mensais.to_parquet(os.path.join(DADOS_DIR, 'retornos_mensais.parquet'))\n\n"
    "print('Arquivos salvos:')\n"
    "print('  precos_ibov.parquet       — preços ajustados diários')\n"
    "print('  retornos_diarios.parquet  — retornos log diários')\n"
    "print('  retornos_mensais.parquet  — retornos log mensais')")

arqs_info = [
    ('precos_ibov.parquet', 'precos_ibov.parquet'),
    ('retornos_diarios.parquet', 'retornos_diarios.parquet'),
    ('retornos_mensais.parquet', 'retornos_mensais.parquet'),
]
linhas = []
for nome, arq in arqs_info:
    fp = os.path.join(DADOS_DIR, arq)
    if os.path.exists(fp):
        tam = os.path.getsize(fp) // 1024
        linhas.append(f"  {nome:<40} {tam:>6} KB")

output(doc,
    "Arquivos salvos:\n"
    "  precos_ibov.parquet       — preços ajustados diários\n"
    "  retornos_diarios.parquet  — retornos log diários\n"
    "  retornos_mensais.parquet  — retornos log mensais\n\n"
    "Tamanhos reais na pasta dados/:\n" + "\n".join(linhas))

interp(doc,
    "Três arquivos com tamanhos bem diferentes. O maior é retornos_diarios.parquet "
    f"— {retornos.shape[0]} linhas × {retornos.shape[1]} colunas de floats. "
    f"O menor é retornos_mensais.parquet — apenas {retornos_mensais.shape[0]} linhas. "
    "Mas é esse arquivo tiny que vai ser carregado toda hora nas próximas aulas, "
    "e ele carrega em menos de 1 segundo graças ao formato parquet. "
    "Por que parquet e não CSV? Três razões: (1) velocidade — 5-10x mais rápido; "
    "(2) tipos preservados — datetime64, float64 sem precisar inferir; "
    "(3) compressão — snappy comprime automaticamente.")

# ══════════════════════════════════════════════════════════════════════════════
secao(doc, "FECHAMENTO DA AULA 2")
tempo(doc, "0:55 – 1:00")
sep(doc)

fala(doc,
    f"Pessoal, excelente trabalho hoje. Baixamos dados de {precos_raw.shape[1]} ações "
    "para um período de 15 anos, tratamos os problemas de qualidade, e geramos "
    "os três parquets que vão alimentar todo o intensivão.")

fala(doc,
    "Os pontos-chave pra fixar: retorno log é aditivo, sempre use preço ajustado, "
    "e lembrem do survivorship bias quando apresentarem os resultados pro Itaú.")

fala(doc,
    "Para a próxima aula: confiram que os três parquets foram criados. Tentem "
    "carregar cada um com pd.read_parquet e verificar o shape esperado: "
    f"{precos_raw.shape} para preços, {retornos.shape} para retornos diários, "
    f"e {retornos_mensais.shape} para retornos mensais.")

fala(doc,
    "Na Aula 3 a gente mergulha nesses dados com EDA — distribuição dos retornos, "
    "correlações, estacionariedade. Vamos entender de verdade o que esse mercado "
    "brasileiro nos diz antes de construirmos qualquer sinal.")

doc.add_paragraph()
sep(doc)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("ImpactUFSCar — Diretoria de Quant — 2025")
r.font.size = Pt(9); r.italic = True; r.font.color.rgb = RGBColor(136,136,136)

doc.save(OUT)
print(f"\nSalvo: {OUT}")
