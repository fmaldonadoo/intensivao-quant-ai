"""
Gera roteiros DOCX para Aulas 8, 9, 10 do Intensivao Quant AI.
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY  = RGBColor(0x0D, 0x1B, 0x3E)
GOLD  = RGBColor(0xF5, 0xA6, 0x23)
GREEN = RGBColor(0x1E, 0x8B, 0x4C)
BLUE  = RGBColor(0x1A, 0x6E, 0xAE)
BROWN = RGBColor(0x7B, 0x4F, 0x12)

BASE = os.path.dirname(os.path.abspath(__file__))


def novo_doc(num, titulo):
    doc = Document()
    sec = doc.sections[0]
    sec.page_width  = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin    = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin   = Inches(1.2)
    sec.right_margin  = Inches(1.2)

    # Cabecalho
    hdr = doc.add_heading("", level=0)
    hdr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = hdr.add_run(f"INTENSIVAO QUANT AI — AULA {num:02d}")
    run.font.color.rgb = NAVY
    run.font.size = Pt(20)
    run.font.bold = True

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run(titulo.upper())
    r.font.color.rgb = GOLD
    r.font.size = Pt(14)
    r.font.bold = True

    doc.add_paragraph()
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rn = note.add_run("ROTEIRO COMPLETO — USO INTERNO DO INSTRUTOR")
    rn.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    rn.font.size = Pt(9)
    rn.font.italic = True
    doc.add_paragraph()
    return doc


def h1(doc, t):
    p = doc.add_heading("", level=1)
    r = p.add_run(t)
    r.font.color.rgb = NAVY
    r.font.size = Pt(15)
    r.font.bold = True
    return p


def h2(doc, t):
    p = doc.add_heading("", level=2)
    r = p.add_run(t)
    r.font.color.rgb = BLUE
    r.font.size = Pt(12)
    r.font.bold = True
    return p


def fala(doc, t):
    p = doc.add_paragraph()
    r = p.add_run(t)
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    p.paragraph_format.space_after = Pt(6)
    return p


def acao(doc, t):
    p = doc.add_paragraph()
    r = p.add_run(f"[{t}]")
    r.font.size = Pt(10)
    r.font.italic = True
    r.font.color.rgb = BROWN
    p.paragraph_format.space_after = Pt(4)
    return p


def cod(doc, t):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(t)
    r.font.name = "Courier New"
    r.font.size = Pt(9)
    r.font.color.rgb = BLUE
    return p


def tempo(doc, t):
    p = doc.add_paragraph()
    r = p.add_run(f"[TEMPO: {t}]")
    r.font.size = Pt(10)
    r.font.bold = True
    r.font.color.rgb = GOLD
    p.paragraph_format.space_after = Pt(2)
    return p


def div(doc):
    p = doc.add_paragraph("─" * 60)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    r = p.runs[0]
    r.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    r.font.size = Pt(8)


def aula08():
    doc = novo_doc(8, "Backtest Rigoroso e Validacao Estatistica")

    h1(doc, "PARTE 1 — TEORIA: O QUE E UM BACKTEST RIGOROSO?")
    tempo(doc, "0:00 – 0:20")
    div(doc)

    h2(doc, "1.1 — O Problema do Backtesting Ingênuo")
    acao(doc, "Slide 1 na tela. Tom sério. Esta aula é a mais importante do curso em termos de rigor científico.")
    fala(doc, "Pessoal, sejam bem-vindos à Aula 8 — Backtest Rigoroso e Validação Estatística.")
    fala(doc, "Nas últimas semanas vocês construíram um pipeline completo: dados, sinal, portfólio, métricas. O backtest mostrou um Sharpe de 0,8, retorno acumulado de mais de 100%. Parece ótimo.")
    fala(doc, "Mas agora vou fazer uma pergunta incômoda: o quanto desse resultado é real, e o quanto é artefato do processo de análise?")
    fala(doc, "Existe um fenômeno muito documentado em finanças quantitativas chamado overfitting — ou ajuste excessivo. Quando você testa muitas hipóteses, muitos parâmetros, muitas variações de um modelo, eventualmente você encontra algo que funcionou no passado simplesmente por acaso estatístico.")
    fala(doc, "Marcos López de Prado, no livro Advances in Financial Machine Learning de 2018, introduziu o Deflated Sharpe Ratio, ou DSR: uma versão do Sharpe que penaliza pela quantidade de estratégias testadas e pela não-normalidade dos retornos. Hoje vamos formalizar isso e aplicar a nossa carteira.")

    h2(doc, "1.2 — As Três Fontes de Viés em Backtests")
    acao(doc, "Slide 2. Enumerate os três vieses, pausando em cada um.")
    fala(doc, "Existem três grandes fontes de viés que precisamos entender e controlar.")
    fala(doc, "Primeiro: Look-ahead bias, ou viés de antecipação. Ocorre quando você usa informação que não estaria disponível no momento da decisão de investimento. No nosso código, controlamos isso com shift(1) no rebalanceamento do backtest.")
    fala(doc, "Segundo: Survivorship bias, ou viés de sobrevivência. Ocorre quando você testa sua estratégia somente em ativos que ainda existem hoje. Vocês devem mencionar essa limitação em seus relatórios finais.")
    fala(doc, "Terceiro: Multiple testing bias, ou viés de múltiplos testes. Se você testa 100 variações de parâmetros, 5 vão parecer excepcionais por mero acaso estatístico. A walk-forward analysis nos protege disso.")

    h2(doc, "1.3 — Walk-Forward Analysis")
    acao(doc, "Slide 3. Diagrama temporal na tela.")
    fala(doc, "A metodologia padrão para mitigar look-back bias e múltiplos testes e o Walk-Forward Analysis.")
    fala(doc, "A ideia e simples: você divide a série temporal em janelas. Em cada janela, você tem um período de treinamento — in-sample — e um período de teste — out-of-sample. Você treina no in-sample, testa no out-of-sample, avança a janela, e repete. Isso garante que cada retorno foi gerado de forma puramente out-of-sample.")

    h2(doc, "1.4 — Deflated Sharpe Ratio (DSR)")
    acao(doc, "Slide 4. Equação do DSR na tela.")
    fala(doc, "O DSR corrige o Sharpe Ratio observado com base em dois fatores: a quantidade de tentativas (múltiplos testes) e a não-normalidade dos retornos (skewness e curtose). Hoje implementaremos uma função robusta para computar essa estatística.")

    h1(doc, "PARTE 2 — CODIGO: BACKTEST_WALKFORWARD E VALIDACAO")
    tempo(doc, "0:20 – 0:55")
    div(doc)

    h2(doc, "2.1 — Setup Inicial")
    acao(doc, "Abrir Jupyter. Novo notebook: aula_08_backtest_rigoroso.ipynb.")
    cod(doc, "import pandas as pd")
    cod(doc, "import numpy as np")
    cod(doc, "import matplotlib.pyplot as plt")
    cod(doc, "from scipy import stats")
    cod(doc, "")
    cod(doc, "ret_mensais   = pd.read_parquet('../dados/retornos_mensais_limpo.parquet')")
    cod(doc, "sinal_v2      = pd.read_parquet('../dados/sinal_v2.parquet')")
    cod(doc, "pesos_v2      = pd.read_parquet('../dados/pesos_v2.parquet')")
    cod(doc, "ret_carteira  = pd.read_parquet('../dados/retorno_carteira_sinal_v2.parquet')")

    h2(doc, "2.2 — Função calcular_dsr()")
    cod(doc, "def calcular_dsr(retornos, n_estrategias=5):")
    cod(doc, "    n  = len(retornos)")
    cod(doc, "    sr = retornos.mean() / retornos.std() * np.sqrt(12)")
    cod(doc, "    sk = retornos.skew()")
    cod(doc, "    ku = retornos.kurtosis()")
    cod(doc, "    e_maxsr = np.sqrt(2) * stats.norm.ppf(1 - 1/n_estrategias) if n_estrategias > 1 else 0")
    cod(doc, "    sr_corr = sr * (1 - sk/6 * sr + (ku-3)/24 * sr**2)**(-0.5)")
    cod(doc, "    dsr = stats.norm.cdf((sr_corr - e_maxsr) * np.sqrt(n - 1) /")
    cod(doc, "                         np.sqrt(1 - sr_corr * e_maxsr + sr_corr**2 * (ku-1)/4))")
    cod(doc, "    return {'sharpe_obs': sr, 'sharpe_corr': sr_corr, 'dsr': dsr}")

    h2(doc, "2.3 — Função backtest_walkforward()")
    cod(doc, "def backtest_walkforward(sinal, ret_mensais, janela_treino=36, janela_teste=6):")
    cod(doc, "    datas = sinal.index.sort_values()")
    cod(doc, "    resultados = []")
    cod(doc, "    inicio = janela_treino")
    cod(doc, "    while inicio + janela_teste <= len(datas):")
    cod(doc, "        datas_oos = datas[inicio:inicio + janela_teste]")
    cod(doc, "        for data in datas_oos:")
    cod(doc, "            if data in sinal.index:")
    cod(doc, "                sig = sinal.loc[data].dropna()")
    cod(doc, "                if len(sig) > 0:")
    cod(doc, "                    ranks  = sig.rank(ascending=False)")
    cod(doc, "                    n_long = max(int(len(ranks) * 0.2), 1)")
    cod(doc, "                    top    = ranks[ranks <= n_long].index")
    cod(doc, "                    pesos  = pd.Series(1/len(top), index=top)")
    cod(doc, "                    if data in ret_mensais.index:")
    cod(doc, "                        ret_mes = ret_mensais.loc[data, pesos.index].dropna()")
    cod(doc, "                        pesos_  = pesos.reindex(ret_mes.index).dropna()")
    cod(doc, "                        pesos_ /= pesos_.sum()")
    cod(doc, "                        ret_oos = (ret_mes * pesos_).sum()")
    cod(doc, "                        resultados.append({'data': data, 'retorno': ret_oos})")
    cod(doc, "        inicio += janela_teste")
    cod(doc, "    return pd.DataFrame(resultados).set_index('data')['retorno']")

    h2(doc, "2.4 — Executar Walk-Forward e Salvar")
    cod(doc, "ret_oos = backtest_walkforward(sinal_v2, ret_mensais, janela_treino=36, janela_teste=6)")
    cod(doc, "dsr_res = calcular_dsr(ret_oos, n_estrategias=5)")
    cod(doc, 'print(f"Sharpe OOS: {dsr_res[\'sharpe_obs\']:.2f}, DSR: {dsr_res[\'dsr\']:.1%}")')
    cod(doc, "ret_oos.to_frame('retorno').to_parquet('../dados/retorno_walkforward_liquido.parquet')")

    div(doc)
    h2(doc, "2.5 — Recap e Proxima Aula")
    tempo(doc, "0:55 – 1:00")
    fala(doc, "Hoje voces aprenderam a diferenca entre um backtest ingenuo e um rigoroso. Implementamos Walk-Forward e calculamos o Deflated Sharpe Ratio (DSR) para medir robustez estatistica.")
    fala(doc, "Na proxima e ultima aula — Aula 9 —, vamos integrar GenAI (Claude API) ao pipeline para gerar narrativas, estruturar o relatorio tecnico vencedor de acordo com os 7 criterios do Itau, e preparar a apresentacao do pitch oral de defesa. Ate semana que vem!")

    folder = os.path.join(BASE, "aula-08-backtest-rigoroso")
    os.makedirs(folder, exist_ok=True)
    doc.save(os.path.join(folder, "roteiro-aula-08-backtest-rigoroso.docx"))
    print("  Aula 08: roteiro salvo -> aula-08-backtest-rigoroso/")


def aula09():
    doc = novo_doc(9, "GenAI, Relatorio Tecnico & Defesa")

    h1(doc, "PARTE 1 — TEORIA: GENAI E OS CRITERIOS DO ITAU")
    tempo(doc, "0:00 – 0:20")
    div(doc)

    h2(doc, "1.1 — LLMs e Prompt Engineering no Pipeline Quant")
    fala(doc, "Bem-vindos à Aula 09 — nossa aula final de consolidacao e entrega! Hoje faremos duas coisas extraordinarias: integraremos a API da Anthropic para gerar a analise de performance automatica da nossa carteira e aprenderemos a estruturar o relatorio técnico e o pitch oral para encantar a banca do Itau.")
    fala(doc, "No mercado moderno, Large Language Models sao usados programaticamente para automatizar a redacao de relatorios de performance, analise de sentimento e auditoria de codigo.")
    fala(doc, "Para obter uma analise premium, aplicamos Prompt Engineering: fornecemos as metricas exatas da nossa carteira (CAGR, Sharpe, Drawdown) e pedimos que o modelo escreva um comentario no estilo de gestora institucional, com rigor tecnico.")

    h2(doc, "1.2 — Os 7 Criterios do Itau e Estrutura da Defesa")
    fala(doc, "A banca do Itau Asset Management vai avaliar o nosso projeto com base em 7 criterios de avaliacao muito especificos:")
    fala(doc, "1. Raciocinio Economico: A fundamentacao do fator momentum (Jegadeesh & Titman 1993, vieses comportamentais).")
    fala(doc, "2. Higiene do Pipeline: Limpeza de outliers, tratamento de gaps e controle do look-ahead bias.")
    fala(doc, "3. Otimizacao da Carteira: A justificativa do uso de Markowitz Restrito para reduzir concentracao e instabilidade amostral.")
    fala(doc, "4. Rigor do Backtest: O uso de Walk-Forward out-of-sample e a deflacao de Sharpe com o DSR.")
    fala(doc, "5. Realismo de Friccoes: O impacto dos custos de transacao e giro na performance liquida.")
    fala(doc, "6. Auto-Critica Honesta: Reconhecer vieses remanescentes, como survivorship bias e regimes de mercado.")
    fala(doc, "7. Qualidade da Entrega: O design premium do relatorio final e a clareza na exposicao oral.")

    h1(doc, "PARTE 2 — CODIGO: CHAMADA DA API E CONTEXTUALIZACAO")
    tempo(doc, "0:20 – 0:55")
    div(doc)

    h2(doc, "2.1 — Setup da API do Claude")
    acao(doc, "Abrir notebook: aulas/aula-09-genai-relatorio/aula-09-genai-relatorio.ipynb")
    cod(doc, "import anthropic")
    cod(doc, "import pandas as pd, numpy as np")
    cod(doc, "")
    cod(doc, "# Inicializa cliente usando variavel de ambiente ANTHROPIC_API_KEY")
    cod(doc, "client = anthropic.Anthropic()")
    cod(doc, "")
    cod(doc, "ret_oos = pd.read_parquet('../dados/retorno_walkforward_liquido.parquet').squeeze()")
    cod(doc, "ret_is  = pd.read_parquet('../dados/retorno_carteira_sinal_v2.parquet').squeeze()")

    h2(doc, "2.2 — Geração de Relatório de Performance")
    cod(doc, "def gerar_comentario_performance(ret_is, ret_oos):")
    cod(doc, "    # Computa metricas")
    cod(doc, "    sh_is = ret_is.mean() / ret_is.std() * np.sqrt(12)")
    cod(doc, "    sh_oos = ret_oos.mean() / ret_oos.std() * np.sqrt(12)")
    cod(doc, "    cagr_oos = (1 + ret_oos).prod() ** (12/len(ret_oos)) - 1")
    cod(doc, "    mdd_oos = ((1 + ret_oos).cumprod() / (1 + ret_oos).cumprod().cummax() - 1).min()")
    cod(doc, "    ")
    cod(doc, "    system = 'Voce e um analista quant senior. Escreva em portugues formal e tecnico, citando os dados fornecidos.'")
    cod(doc, "    prompt = f'''Analise a performance da nossa carteira de Momentum v2:")
    cod(doc, "    - In-Sample Sharpe: {sh_is:.2f}")
    cod(doc, "    - Out-of-Sample Sharpe: {sh_oos:.2f}")
    cod(doc, "    - CAGR OOS: {cagr_oos:.1%}")
    cod(doc, "    - Max Drawdown OOS: {mdd_oos:.1%}")
    cod(doc, "    Escreva um relatorio analitico e objetivo de ate 200 palavras para a banca do Itau.'''")
    cod(doc, "    ")
    cod(doc, "    msg = client.messages.create(")
    cod(doc, "        model='claude-3-5-haiku-20241022',")
    cod(doc, "        max_tokens=1024,")
    cod(doc, "        temperature=0.0,")
    cod(doc, "        system=system,")
    cod(doc, "        messages=[{'role': 'user', 'content': prompt}]")
    cod(doc, "    )")
    cod(doc, "    return msg.content[0].text")
    cod(doc, "")
    cod(doc, "relatorio = gerar_comentario_performance(ret_is, ret_oos)")
    cod(doc, "print(relatorio)")

    div(doc)
    h2(doc, "2.3 — Simulação e Estrutura da Apresentação Oral")
    tempo(doc, "0:55 – 1:00")
    fala(doc, "Excelente. O modelo gera a narrativa ideal para integrarmos no capitulo final do nosso relatorio. Agora, os minutos finais sao para preparar voces para o Pitch Oral.")
    fala(doc, "Estrutura sugerida do Pitch de 10 minutos: 2 min para Tese/Dados, 3 min para a Metodologia (Rigor, Markowitz Restrito), 3 min para Resultados Líquidos e Sensibilidade, e 2 min para a Auto-Critica (limitações) e Encerramento.")
    fala(doc, "Dica de ouro para as perguntas da banca: se perguntarem por que nao usaram Markowitz Puro, respondam com base no que vimos na Aula 6: 'Markowitz Puro tem turnover excessivo (X%) e alta concentracao de pesos por conta de erros de estimativa amostral; a versao Restrita (bounds de 20%) mitigou isso com robustez out-of-sample.'")
    fala(doc, "Pessoal, encerramos aqui o Intensivao Quant AI! Voces construiram um pipeline de nivel profissional, do zero absoluto ao relatorio final integrado com GenAI. Muito obrigado pela dedicação e boa sorte no Desafio Itaú! Voces estao prontos para vencer!")

    folder = os.path.join(BASE, "aula-09-genai-relatorio")
    os.makedirs(folder, exist_ok=True)
    doc.save(os.path.join(folder, "roteiro-aula-09-genai-relatorio.docx"))
    print("  Aula 09: roteiro salvo -> aula-09-genai-relatorio/")


if __name__ == "__main__":
    print("Gerando roteiros das aulas 08 e 09...")
    aula08()
    aula09()
    print("Concluido!")
